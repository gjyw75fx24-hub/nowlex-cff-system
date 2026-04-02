# contratos/views.py

from django.http import JsonResponse, HttpResponse, FileResponse
from django.shortcuts import render, get_object_or_404
from django.views.decorators.http import require_POST, require_GET
from .models import (
    ProcessoJudicial, ProcessoJudicialNumeroCnj, Parte, StatusProcessual, QuestaoAnalise,
    OpcaoResposta, Contrato, ProcessoArquivo, DocumentoModelo, TipoAnaliseObjetiva
)
from .permissoes import filter_processos_queryset_for_user
from .integracoes_escavador.api import buscar_processo_por_cnj
from .integracoes_escavador.partes import collect_partes_from_escavador_payload
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP, ROUND_CEILING
from django.db.models import Max
from django.db import transaction
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth.decorators import login_required
from django.urls import reverse
import re
import logging
import requests
from urllib.parse import quote
from contratos.data.decision_tree_config import DECISION_TREE_CONFIG # <--- Nova importação
import copy # Para cópia profunda do dicionário
import json
import time
import tempfile
from pathlib import Path
import subprocess
import threading
import shutil
import zipfile

# Imports para geração de DOCX
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor
from io import BytesIO
from datetime import datetime
import os
from django.conf import settings
from django.core.files.base import ContentFile
from django.utils.text import slugify
from django.utils.timezone import now
from django.db.models import Q
from django.db.models.query import QuerySet
from django.db.utils import OperationalError, ProgrammingError


logger = logging.getLogger(__name__)


def _sanitize_filename(name: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]+', ' ', name or '')
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned or 'arquivo'


def _format_currency_brl(value):
    try:
        amount = Decimal(value or 0)
    except (TypeError, InvalidOperation):
        amount = Decimal('0')
    quantized = amount.quantize(Decimal('0.01'))
    formatted = f'{quantized:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.')
    return f'R$ {formatted}'


def _to_decimal(value):
    if isinstance(value, Decimal):
        return value
    if value is None:
        return Decimal('0')
    try:
        return Decimal(str(value))
    except (InvalidOperation, ValueError):
        return Decimal('0')


def _format_cpf(cpf_value):
    digits = re.sub(r'\D', '', str(cpf_value or ''))
    if len(digits) == 11:
        return f'{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:]}'
    return str(cpf_value or '')


def _parse_decimal_input(raw):
    if raw is None:
        return None
    text = str(raw).strip()
    if not text:
        return None
    text = text.replace('R$', '').replace(' ', '')
    if ',' in text and '.' in text:
        text = text.replace('.', '').replace(',', '.')
    elif ',' in text:
        text = text.replace(',', '.')
    try:
        return Decimal(text)
    except (InvalidOperation, ValueError):
        return None


def _parse_int_input(raw):
    try:
        value = int(str(raw).strip())
    except (TypeError, ValueError):
        return None
    return value if value > 0 else None


def _build_endereco_from_parts(parts):
    if not isinstance(parts, dict):
        return ''
    fields = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
    chunks = []
    for field in fields:
        value = str(parts.get(field, '') or '').strip()
        chunks.append(f'{field}: {value}')
    return ' - '.join(chunks).strip()


def _extrair_primeiros_nomes(nome, max_nomes=2):
    if not nome:
        return ''
    STOP = {'da', 'de', 'do', 'das', 'dos', 'e', "d’", "d'", 'a', 'o'}
    tokens = [t for t in re.split(r'\s+', nome.strip()) if t]
    out = []
    for tok in tokens:
        if tok.lower() in STOP:
            continue
        out.append(tok)
        if len(out) >= max_nomes:
            break
    return ' '.join(out).strip()


def _formatar_lista_contratos(contratos):
    valores = [str(c.numero_contrato).strip() for c in contratos if getattr(c, 'numero_contrato', None)]
    if not valores:
        return ''
    if len(valores) == 1:
        return valores[0]
    ultimo = valores.pop()
    return f"{', '.join(valores)} e {ultimo}"


def _normalize_digits(value):
    return re.sub(r'\D', '', str(value or ''))


def _parse_contract_ids(values):
    parsed = []
    seen = set()
    for item in values or []:
        if isinstance(item, int):
            contract_id = item
        else:
            digits = _normalize_digits(item)
            if not digits:
                continue
            try:
                contract_id = int(digits)
            except ValueError:
                continue
        if contract_id not in seen:
            seen.add(contract_id)
            parsed.append(contract_id)
    return parsed


def _sanitize_contract_numbers(contracts):
    seen = []
    for contract in contracts:
        number = getattr(contract, 'numero_contrato', None)
        digits = _normalize_digits(number)
        if digits and digits not in seen:
            seen.append(digits)
    return seen


def _format_contracts_label(numbers):
    if not numbers:
        return ''
    if len(numbers) == 1:
        return numbers[0]
    last = numbers[-1]
    return f"{', '.join(numbers[:-1])} e {last}"


def _build_extrato_filename(contracts_label, parte_name, cpf_digits):
    label_part = contracts_label or 'contratos'
    nomes = _extrair_primeiros_nomes(parte_name)
    if not nomes:
        nomes = cpf_digits or 'perfil'
    filename = f"05 - Extrato de Titularidade - {label_part} - {nomes}.pdf"
    return _sanitize_filename(filename)


def _call_nowlex_extrato(cpf_digits, contract_numbers):
    api_key = getattr(settings, 'NOWLEX_JUDICIAL_API_KEY', '') or os.environ.get('NOWLEX_JUDICIAL_API_KEY', '')
    if not api_key:
        raise RuntimeError('Chave NOWLEX_JUDICIAL_API_KEY não configurada.')
    included = ','.join(contract_numbers)
    params = quote(included, safe=',')
    url = f"https://erp-api.nowlex.com/api/judicial/cpf/{cpf_digits}/pdf?include_contracts={params}"
    response = requests.get(url, headers={'X-API-Key': api_key}, timeout=60)
    if not response.ok:
        raise RuntimeError(f"Status {response.status_code}: {response.text}")
    if not response.content:
        raise RuntimeError("Resposta da API sem conteúdo.")
    return response.content


def generate_extrato_titularidade(processo, cpf_value, contratos, parte_name, usuario):
    cpf_digits = _normalize_digits(cpf_value)
    if len(cpf_digits) != 11:
        return {'ok': False, 'error': 'CPF inválido para o extrato.'}
    contract_numbers = _sanitize_contract_numbers(contratos)
    if not contract_numbers:
        return {'ok': False, 'error': 'Nenhum contrato válido para o extrato.'}
    contratos_label = _format_contracts_label(contract_numbers)
    try:
        pdf_bytes = _call_nowlex_extrato(cpf_digits, contract_numbers)
    except Exception as exc:
        logger.error("Falha ao gerar extrato de titularidade: %s", exc, exc_info=True)
        error_message = str(exc)
        if "Status 404" in error_message or "Nenhum contrato" in error_message:
            error_message = (
                "NowLex não possui o cadastro do contrato solicitado. "
                "Não é possível emitir o extrato da titularidade."
            )
        if 'NOWLEX_JUDICIAL_API_KEY' in error_message:
            # Permite seguir sem o extrato se a chave não estiver configurada
            return {'ok': False, 'error': 'Extrato não gerado (chave NowLex ausente).'}
        return {'ok': False, 'error': error_message}
    filename = _build_extrato_filename(contratos_label, parte_name, cpf_digits)
    pdf_file = ContentFile(pdf_bytes)
    arquivo = ProcessoArquivo(
        processo=processo,
        nome=filename,
        enviado_por=usuario if usuario and usuario.is_authenticated else None,
    )
    arquivo.arquivo.save(filename, pdf_file, save=True)
    return {'ok': True, 'pdf_url': arquivo.arquivo.url}


def _iter_container_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in getattr(container, 'tables', []):
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_container_paragraphs(cell)


def _bold_keywords_in_document(document, keywords):
    upper_keywords = [kw.upper() for kw in keywords]
    for paragraph in _iter_container_paragraphs(document):
        for run in paragraph.runs:
            run_text_upper = run.text.upper()
            if any(kw in run_text_upper for kw in upper_keywords):
                run.font.bold = True
                run.font.name = 'Times New Roman'


def _replace_placeholder_across_runs(paragraph, placeholder, replacement):
    if not placeholder:
        return False
    replacement = str(replacement)
    runs = list(paragraph.runs)
    if not runs:
        return False
    full_text = ''.join(run.text for run in runs)
    if placeholder not in full_text:
        return False

    while True:
        full_text = ''.join(run.text for run in runs)
        start = full_text.find(placeholder)
        if start == -1:
            break
        end = start + len(placeholder)

        current_index = 0
        start_run_idx = None
        start_offset = 0
        end_run_idx = None
        end_offset = 0
        for idx, run in enumerate(runs):
            run_len = len(run.text)
            if start_run_idx is None and current_index + run_len >= start:
                start_run_idx = idx
                start_offset = start - current_index
            if current_index + run_len >= end:
                end_run_idx = idx
                end_offset = end - current_index
                break
            current_index += run_len

        if start_run_idx is None or end_run_idx is None:
            break

        if start_run_idx == end_run_idx:
            run = runs[start_run_idx]
            run.text = (
                run.text[:start_offset]
                + replacement
                + run.text[end_offset:]
            )
        else:
            start_run = runs[start_run_idx]
            end_run = runs[end_run_idx]
            start_run.text = (
                start_run.text[:start_offset]
                + replacement
                + end_run.text[end_offset:]
            )
            for idx in range(start_run_idx + 1, end_run_idx + 1):
                runs[idx].text = ''
    return True


def _replace_placeholders_in_paragraph(paragraph, data):
    if not paragraph or not data:
        return
    if '[E]/[H]' in paragraph.text:
        _replace_placeholder_across_runs(
            paragraph,
            '[E]/[H]',
            f"{data.get('E_FORO', '')}/{data.get('H_FORO', '')}",
        )
    for key, value in data.items():
        _replace_placeholder_across_runs(paragraph, f'[{key}]', value)


def _replace_placeholders_in_container(container, data):
    if not container or not data:
        return
    for paragraph in _iter_container_paragraphs(container):
        _replace_placeholders_in_paragraph(paragraph, data)


def _delete_paragraph(paragraph):
    if paragraph is None or paragraph._element is None:
        return
    parent = paragraph._element.getparent()
    if parent is None:
        return
    parent.remove(paragraph._element)


def _remove_paragraphs_containing(container, text_substring):
    if not container or not text_substring:
        return
    for paragraph in list(_iter_container_paragraphs(container)):
        if text_substring in paragraph.text:
            _delete_paragraph(paragraph)


def _replace_paragraph_text(paragraph, replacement):
    if paragraph is None or replacement is None:
        return
    for run in list(paragraph.runs):
        run.text = ''
    lines = str(replacement).splitlines()
    if not lines:
        return
    for idx, line in enumerate(lines):
        if idx > 0:
            paragraph.add_run().add_break()
        paragraph.add_run(line)


def _replace_paragraphs_containing(container, text_substring, replacement, first_only=False):
    if not container or not text_substring or replacement is None:
        return
    for paragraph in list(_iter_container_paragraphs(container)):
        if text_substring in paragraph.text:
            _replace_paragraph_text(paragraph, replacement)
            if first_only:
                break


def _extract_custas_block_text(document, start_markers, end_markers=None):
    if not document or not start_markers:
        return ''
    collecting = False
    lines = []
    marker_set = tuple(start_markers)
    end_set = tuple(end_markers or [])
    for paragraph in _iter_container_paragraphs(document):
        text = (paragraph.text or '').strip()
        if not text:
            continue
        has_marker = any(marker in text for marker in marker_set)
        if not collecting and has_marker:
            collecting = True
        if collecting:
            if end_set and any(marker in text for marker in end_set) and not has_marker:
                break
            if re.match(r'^[IVXLCDM]+\.', text) and not has_marker:
                break
            lines.append(text)
    return '\n'.join(lines).strip()


def _extract_custas_block_from_docx_bytes(docx_bytes, start_markers, end_markers=None):
    if not docx_bytes:
        return ''
    try:
        document = Document(BytesIO(docx_bytes))
    except Exception:
        return ''
    return _extract_custas_block_text(document, start_markers, end_markers=end_markers)


def _complete_cobranca_custas_preview(text, custas_valor, parcelas_override=None, valor_parcela_override=None):
    if not text or custas_valor is None:
        return text or ''
    parcelas, valor_parcela = _resolve_custas_parcelamento(
        custas_valor,
        parcelas_override=parcelas_override,
        valor_parcela_override=valor_parcela_override,
    )
    custas_formatadas = _format_currency_brl(custas_valor)
    parcela_formatada = _format_currency_brl(valor_parcela) if valor_parcela is not None else ''
    parcelas_extenso = number_to_words_pt_br(
        parcelas,
        feminine=True,
        include_currency=False,
        capitalize_first=False
    ) if parcelas else ''
    lines = []
    for line in text.splitlines():
        trimmed = line.rstrip()
        if 'A requerente requer o parcelamento em' in trimmed and parcelas and parcelas_extenso:
            trimmed = f"A requerente requer o parcelamento em {parcelas} ({parcelas_extenso}) parcelas mensais e"
        if 'milhares de' in trimmed and ('R$' not in trimmed and 'reais' not in trimmed):
            trimmed = re.sub(
                r'(milhares de)\s*$',
                f"\\1 reais ({custas_formatadas})",
                trimmed,
                flags=re.IGNORECASE
            )
        if parcela_formatada and 'parcelas mensais' in trimmed and 'valor de' not in trimmed:
            trimmed = re.sub(
                r'(parcelas mensais(?:\s+e)?)\s*$',
                f"\\1 sucessivas no valor de {parcela_formatada}.",
                trimmed,
                flags=re.IGNORECASE
            )
        lines.append(trimmed)
    return '\n'.join(lines).strip()


def _sanitize_monitoria_custas_paragrafo(text):
    if not text:
        return ''
    sanitized = str(text).replace('\r\n', '\n').replace('\r', '\n')
    stop_markers = (
        'Por fim, requer-se',
        'Dá-se à causa',
        'Nestes termos',
        'DEFERIMENTO',
        'São Paulo,',
    )
    for marker in stop_markers:
        marker_index = sanitized.find(marker)
        if marker_index >= 0:
            sanitized = sanitized[:marker_index]
            break
    lines = [line.strip() for line in sanitized.split('\n')]
    collapsed = []
    last_blank = True
    for line in lines:
        if not line:
            if not last_blank:
                collapsed.append('')
            last_blank = True
            continue
        collapsed.append(line)
        last_blank = False
    return '\n'.join(collapsed).strip()


def _should_replace_monitoria_custas_paragrafo(text):
    normalized = _sanitize_monitoria_custas_paragrafo(text)
    if not normalized:
        return False
    # Frases muito curtas indicam que o front não conseguiu carregar a prévia
    # completa; nesse caso é melhor preservar o parágrafo original do modelo.
    if len(normalized) < 120:
        return False
    return 'parcelamento das custas iniciais' in normalized.lower()


def _complete_cobranca_custas_paragraphs(document, custas_valor, parcelas, valor_parcela):
    if not document or custas_valor is None:
        return
    for paragraph in _iter_container_paragraphs(document):
        raw_text = paragraph.text or ''
        if not raw_text:
            continue
        if not any(marker in raw_text for marker in ('milhares de', 'parcelas mensais', 'A requerente requer o parcelamento em')):
            continue
        updated = _complete_cobranca_custas_preview(
            raw_text,
            custas_valor,
            parcelas_override=parcelas,
            valor_parcela_override=valor_parcela
        )
        if updated and updated != raw_text:
            _replace_paragraph_text(paragraph, updated)


def _replacePlaceholderStyled_(document, pattern, replacement, bold=False):
    if not pattern or replacement is None:
        return
    for paragraph in _iter_container_paragraphs(document):
        paragraph_text = ''.join(run.text for run in paragraph.runs)
        if pattern not in paragraph_text:
            continue
        segments = paragraph_text.split(pattern)
        if len(segments) <= 1:
            continue
        for run in paragraph.runs:
            run.text = ''
        for index, segment in enumerate(segments):
            if segment:
                run = paragraph.add_run(segment)
                run.font.name = 'Times New Roman'
            if index < len(segments) - 1:
                replacement_run = paragraph.add_run(replacement)
                replacement_run.font.name = 'Times New Roman'
                if bold:
                    replacement_run.bold = True
                    replacement_run.font.bold = True


def _bold_paragraphs_containing(document, text):
    if not text:
        return
    normalized = text.strip()
    for paragraph in _iter_container_paragraphs(document):
        if normalized in paragraph.text:
            for run in paragraph.runs:
                run.bold = True
                run.font.bold = True


def _set_run_shading(run, fill_color):
    if not fill_color:
        return
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    rPr.append(shd)


def _parse_placeholder_segments(text):
    if not text:
        return []
    pattern = re.compile(r'\[(lg|n|a|ag|m)\]')
    segments = []
    state = {
        'bold': False,
        'blue_font': False,
        'highlight_blue': False,
        'highlight_line': False,
        'line_highlight_open': False,
        'uppercase': False,
    }
    last_index = 0

    def _snapshot():
        return {
            'bold': state['bold'],
            'blue_font': state['blue_font'],
            'highlight_blue': state['highlight_blue'],
            'highlight_line': state['highlight_line'],
            'uppercase': state['uppercase'],
        }

    for match in pattern.finditer(text):
        if match.start() > last_index:
            segments.append((text[last_index:match.start()], _snapshot()))
        token = match.group(1)
        if token == 'n':
            state['bold'] = not state['bold']
            if not state['bold'] and state['line_highlight_open']:
                state['highlight_line'] = False
                state['line_highlight_open'] = False
        elif token == 'lg':
            state['highlight_line'] = True
            state['line_highlight_open'] = True
        elif token == 'a':
            state['blue_font'] = not state['blue_font']
        elif token == 'ag':
            state['highlight_blue'] = not state['highlight_blue']
        elif token == 'm':
            state['uppercase'] = not state['uppercase']
        last_index = match.end()

    if last_index < len(text):
        segments.append((text[last_index:], _snapshot()))
    return [seg for seg in segments if seg[0]]


def _apply_placeholder_styles(document):
    markers = ['[n]', '[a]', '[ag]', '[lg]', '[m]']
    highlight_fill = 'DCE7FB'
    blue_rgb = RGBColor(0x33, 0x3A, 0xF1)

    for paragraph in _iter_container_paragraphs(document):
        source_text = paragraph.text
        if not source_text or not any(marker in source_text for marker in markers):
            continue
        segments = _parse_placeholder_segments(source_text)
        if not segments:
            continue
        for run in paragraph.runs:
            run.text = ''
        for text_segment, style in segments:
            run_text = text_segment.upper() if style['uppercase'] else text_segment
            run = paragraph.add_run(run_text)
            run.font.name = 'Times New Roman'
            run.font.bold = style['bold']
            run.font.color.rgb = blue_rgb if (style['blue_font'] or style['highlight_blue']) else None
            if style['highlight_line'] or style['highlight_blue']:
                _set_run_shading(run, highlight_fill)


def _load_template_document(slug, fallback_path=None):
    """Carrega template do banco de dados (S3 em produção)."""
    template = DocumentoModelo.objects.filter(slug=slug).first()
    if not template:
        label = dict(DocumentoModelo.SlugChoices.choices).get(slug, '')
        if label:
            template = DocumentoModelo.objects.filter(nome__iexact=label).order_by('-atualizado_em').first()
    if template:
        try:
            with template.arquivo.open('rb') as handle:
                data = BytesIO(handle.read())
            data.seek(0)
            return Document(data)
        except (ValueError, FileNotFoundError) as e:
            raise FileNotFoundError(
                f"Arquivo do template '{slug}' não encontrado no S3. "
                f"Verifique se o arquivo foi enviado corretamente. Erro: {e}"
            )
    raise FileNotFoundError(
        f"Template de documento não encontrado (slug={slug}). "
        "Cadastre o arquivo via Admin > Documentos Modelo."
    )


def _calculate_monitoria_installments(amount, target=Decimal('500'), max_installments=10):
    """Calcula o número de parcelas considerando o valor alvo e limite."""
    if amount <= Decimal('0'):
        return 1
    installments_decimal = (amount / target).to_integral_value(rounding=ROUND_CEILING)
    installments = int(installments_decimal)
    if installments < 1:
        installments = 1
    return min(installments, max_installments)


def _resolve_custas_parcelamento(custas_valor, parcelas_override=None, valor_parcela_override=None):
    if custas_valor is None:
        return None, None
    parcelas = parcelas_override if parcelas_override is not None else _calculate_monitoria_installments(custas_valor)
    try:
        parcelas = int(parcelas)
    except (TypeError, ValueError):
        parcelas = _calculate_monitoria_installments(custas_valor)
    if parcelas <= 0:
        parcelas = _calculate_monitoria_installments(custas_valor)
    if valor_parcela_override is not None:
        valor_parcela = _to_decimal(valor_parcela_override).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    else:
        valor_parcela = (custas_valor / Decimal(parcelas)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    return parcelas, valor_parcela

def _normalize_monitoria_contracts(contracts):
    if isinstance(contracts, QuerySet):
        try:
            contracts = contracts.select_related('processo')
        except AttributeError:
            pass
    if contracts is None:
        return []
    return list(contracts)


def _build_docx_bytes_common(
    processo,
    polo_passivo,
    contratos_monitoria,
    processo_override=None,
    custas_override=None,
    custas_paragrafo=None,
    custas_parcelas=None,
    custas_valor_parcela=None
):
    contratos_monitoria = _normalize_monitoria_contracts(contratos_monitoria)
    processo_override = processo_override or {}
    override_vara = processo_override.get('vara')
    override_uf = processo_override.get('uf')
    override_valor_causa = processo_override.get('valor_causa')
    dados = {}
    parte_nome = (polo_passivo.nome or '').strip()
    dados['PARTE CONTRÁRIA'] = f"[n]{parte_nome}[n]" if parte_nome else ''
    dados['CPF'] = polo_passivo.documento

    endereco_parts = parse_endereco(polo_passivo.endereco)
    dados['A'] = _format_address_component(endereco_parts.get('A', '') or '')
    dados['B'] = _format_address_component(endereco_parts.get('B', '') or '')
    dados['C'] = _format_address_component(endereco_parts.get('C', '') or '')
    dados['D'] = _format_address_component(endereco_parts.get('D', '') or '')
    dados['E'] = _format_address_component(endereco_parts.get('E', '') or '')
    dados['F'] = _format_address_component(endereco_parts.get('F', '') or '')
    dados['G'] = _format_address_component(endereco_parts.get('G', '') or '')
    dados['H'] = _format_address_component(endereco_parts.get('H', '') or '')

    dados['E_FORO'] = override_vara if override_vara is not None else processo.vara
    dados['H_FORO'] = override_uf if override_uf is not None else processo.uf

    dados['CONTRATO'] = ", ".join([c.numero_contrato for c in contratos_monitoria if c.numero_contrato])

    contrato_total = sum(
        (_to_decimal(contrato.valor_causa) for contrato in contratos_monitoria),
        Decimal('0')
    ) if contratos_monitoria else Decimal('0')
    total_valor_causa = contrato_total
    if total_valor_causa == Decimal('0'):
        if override_valor_causa is not None:
            total_valor_causa = _to_decimal(override_valor_causa)
        else:
            total_valor_causa = _to_decimal(processo.valor_causa)

    dados['VALOR DA CAUSA'] = _format_currency_brl(total_valor_causa)
    dados['VALOR DA CAUSA POR EXTENSO'] = number_to_words_pt_br(total_valor_causa)

    custas_rate = Decimal('0.025')
    if custas_override is not None:
        valor_custas_iniciais = _to_decimal(custas_override).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    else:
        valor_custas_iniciais = (total_valor_causa * custas_rate).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    valor_custas_texto = number_to_words_pt_br(valor_custas_iniciais)
    dados['2,5 % DO VALOR DA CAUSA'] = _format_currency_brl(valor_custas_iniciais)
    dados['2,5% DO VALOR DA CAUSA'] = dados['2,5 % DO VALOR DA CAUSA']
    dados['2,5 % DO VALOR DA CAUSA POR EXTENSO'] = valor_custas_texto
    dados['2,5% DO VALOR DA CAUSA POR EXTENSO'] = valor_custas_texto
    parcelas_custas, valor_parcela = _resolve_custas_parcelamento(
        valor_custas_iniciais,
        parcelas_override=custas_parcelas,
        valor_parcela_override=custas_valor_parcela
    )
    dados['X PARCELAS'] = str(parcelas_custas)
    dados['X PARCELAS POR EXTENSO'] = number_to_words_pt_br(
        parcelas_custas,
        feminine=True,
        include_currency=False,
        capitalize_first=False
    )

    dados['DATA DE HOJE'] = datetime.now().strftime("%d de %B de %Y").replace(
        'January', 'janeiro').replace('February', 'fevereiro').replace('March', 'março').replace(
        'April', 'abril').replace('May', 'maio').replace('June', 'junho').replace(
        'July', 'julho').replace('August', 'agosto').replace('September', 'setembro').replace(
        'October', 'outubro').replace('November', 'novembro').replace('December', 'dezembro')

    document = _load_template_document(DocumentoModelo.SlugChoices.MONITORIA_INICIAL, None)

    show_parcelamento = valor_custas_iniciais >= Decimal('1000')

    # Ajusta posição do rodapé para evitar cortes no PDF (mantém margens do template)
    for section in document.sections:
        try:
            # mantém margens originais do template, apenas sobe o rodapé
            section.footer_distance = Cm(1.5)
        except Exception:
            pass

    if not show_parcelamento:
        _remove_paragraphs_containing(
            document,
            "Seja deferido o parcelamento das custas iniciais"
        )
    elif _should_replace_monitoria_custas_paragrafo(custas_paragrafo):
        _replace_paragraphs_containing(
            document,
            "Seja deferido o parcelamento das custas iniciais",
            _sanitize_monitoria_custas_paragrafo(custas_paragrafo),
            first_only=True
        )

    _replace_placeholders_in_container(document, dados)
    for section in document.sections:
        _replace_placeholders_in_container(section.header, dados)
        _replace_placeholders_in_container(section.footer, dados)

    _apply_placeholder_styles(document)
    _bold_keywords_in_document(document, ['EXCELENTÍSSIMO(A)'])

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()


def _build_monitoria_base_filename(polo_passivo, contratos_monitoria):
    contratos_labels = [
        (c.numero_contrato or f"contrato-{c.id}") for c in contratos_monitoria
    ]
    if not contratos_labels:
        contratos_labels = ['contratos']
    contratos_segment = " - ".join(contratos_labels)
    nome_parte = polo_passivo.nome or "parte"
    base = f"01 - Monitoria Inicial - {contratos_segment} - {nome_parte}"
    return _sanitize_filename(base)


def _get_total_contrato_value(contratos, processo, processo_valor_causa=None):
    total = Decimal('0')
    for contrato in contratos:
        valor = contrato.valor_total_devido if contrato.valor_total_devido is not None else contrato.valor_causa
        if valor is not None:
            total += valor
    if total == Decimal('0'):
        if processo_valor_causa is not None:
            total = processo_valor_causa
        elif processo.valor_causa is not None:
            total = processo.valor_causa
    return total


def _get_total_valor_causa(contratos, processo, processo_valor_causa=None):
    total = Decimal('0')
    for contrato in contratos:
        if contrato.valor_causa is not None:
            total += _to_decimal(contrato.valor_causa)
    if total == Decimal('0'):
        if processo_valor_causa is not None:
            total = _to_decimal(processo_valor_causa)
        else:
            total = _to_decimal(processo.valor_causa)
    return total


def _build_cobranca_base_filename(polo_passivo, contratos):
    label = _formatar_lista_contratos(contratos) or 'contratos'
    nome_parte = _extrair_primeiros_nomes(polo_passivo.nome or '', 2) or 'parte'
    base = f"01 - Cobranca Judicial - {label} - {nome_parte}"
    return _sanitize_filename(base)


def _build_habilitacao_base_filename(polo_passivo, processo, cnj_reference=None):
    nome_parte = _extrair_primeiros_nomes(polo_passivo.nome or '', 2) or 'parte'
    reference = cnj_reference or processo.cnj or f'processo-{processo.pk}'
    base = f"Habilitacao - {nome_parte} - {reference}"
    return _sanitize_filename(base)


def _convert_docx_to_pdf_bytes(docx_bytes, *, allow_remote=True, gotenberg_timeout=120):
    """
    Converte DOCX para PDF.
    Prioriza Gotenberg (serviço com LibreOffice embutido) quando permitido.
    Se não disponível, tenta LibreOffice local (soffice/libreoffice).
    Como fallback, usa mammoth + xhtml2pdf (100% Python).
    Último recurso: reportlab direto do DOCX (texto/tabelas).
    """
    gotenberg_url = ''
    if allow_remote:
        gotenberg_url = getattr(settings, 'GOTENBERG_URL', os.environ.get('GOTENBERG_URL', ''))

    def _convert_with_gotenberg():
        if not gotenberg_url:
            return None
        try:
            logger.info("Tentando conversão via Gotenberg: %s", gotenberg_url)
            files = {'files': ('document.docx', docx_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = requests.post(
                f"{gotenberg_url}/forms/libreoffice/convert",
                files=files,
                timeout=gotenberg_timeout
            )
            if response.status_code == 200 and response.content:
                pdf_size = len(response.content)
                logger.info("Gotenberg: conversão bem-sucedida (PDF: %d bytes)", pdf_size)

                if response.content[:5] == b'%PDF-':
                    if pdf_size > 1000:
                        return response.content
                    logger.warning("Gotenberg: PDF muito pequeno (%d bytes), pode estar corrompido", pdf_size)
                else:
                    logger.warning("Gotenberg: Conteúdo retornado não é um PDF válido")
            else:
                logger.warning("Gotenberg falhou: status=%s", response.status_code)
        except requests.Timeout:
            logger.warning("Gotenberg timeout após %ss", gotenberg_timeout)
        except Exception as exc:
            logger.warning("Erro ao usar Gotenberg: %s", exc)
        return None

    if gotenberg_url:
        gotenberg_pdf = _convert_with_gotenberg()
        if gotenberg_pdf:
            return gotenberg_pdf

    def _find_soffice():
        # Lista expandida de localizações possíveis do LibreOffice
        candidates = [
            "soffice",
            "/usr/bin/soffice",
            "libreoffice",
            "/usr/bin/libreoffice",
            "/opt/libreoffice/program/soffice",
            "/usr/lib/libreoffice/program/soffice",
            "/snap/bin/libreoffice",
        ]
        for candidate in candidates:
            if shutil.which(candidate):
                logger.info("LibreOffice encontrado em: %s", candidate)
                return candidate
        # Log de diagnóstico se não encontrar
        logger.error("LibreOffice não encontrado. Tentou: %s", ", ".join(candidates))
        logger.error("PATH atual: %s", os.environ.get('PATH', 'N/A'))
        return None

    # Tenta LibreOffice local
    soffice_cmd = _find_soffice()
    if soffice_cmd:
        logger.info("LibreOffice encontrado em: %s", soffice_cmd)
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                tmpdir_path = Path(tmpdir)
                docx_path = tmpdir_path / "input.docx"
                pdf_path = tmpdir_path / "input.pdf"
                docx_path.write_bytes(docx_bytes)
                cmd = [
                    soffice_cmd,
                    "--headless",
                    "--nologo",
                    "--nodefault",
                    "--norestore",
                    "--nofirststartwizard",
                    "--convert-to",
                    "pdf:writer_pdf_Export",
                    "--outdir",
                    str(tmpdir_path),
                    str(docx_path),
                ]
                logger.info("Executando LibreOffice: %s", ' '.join(cmd))
                result = subprocess.run(cmd, capture_output=True, timeout=90)
                if result.returncode == 0 and pdf_path.exists():
                    logger.info("LibreOffice: conversão bem-sucedida")
                    return pdf_path.read_bytes()
                logger.warning(
                    "LibreOffice falhou, tentando fallback: rc=%s stdout=%s stderr=%s",
                    result.returncode,
                    result.stdout.decode('utf-8', errors='ignore')[:200],
                    result.stderr.decode('utf-8', errors='ignore')[:200],
                )
        except Exception as exc:
            logger.warning("Erro com LibreOffice, tentando fallback: %s", exc, exc_info=True)
    else:
        logger.warning("LibreOffice não encontrado, usando fallback Python")

    # Fallback: mammoth + xhtml2pdf (100% Python, sem dependências externas)
    try:
        import mammoth
        from xhtml2pdf import pisa

        # Converte DOCX para HTML usando mammoth
        docx_io = BytesIO(docx_bytes)
        result = mammoth.convert_to_html(docx_io)
        html_content = result.value

        # Adiciona CSS para melhor formatação do PDF
        html_with_style = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                body {{
                    font-family: 'Times New Roman', Times, serif;
                    font-size: 12pt;
                    line-height: 1.5;
                    color: #000;
                }}
                p {{
                    margin: 0 0 10pt 0;
                    text-align: justify;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 10pt 0;
                }}
                td, th {{
                    border: 1px solid #000;
                    padding: 5pt;
                }}
                h1 {{ font-size: 14pt; font-weight: bold; margin: 12pt 0 6pt 0; }}
                h2 {{ font-size: 13pt; font-weight: bold; margin: 10pt 0 5pt 0; }}
                h3 {{ font-size: 12pt; font-weight: bold; margin: 8pt 0 4pt 0; }}
                strong, b {{ font-weight: bold; }}
                em, i {{ font-style: italic; }}
                u {{ text-decoration: underline; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        # Converte HTML para PDF usando xhtml2pdf
        pdf_io = BytesIO()
        pisa_status = pisa.CreatePDF(html_with_style, dest=pdf_io, encoding='UTF-8')

        if pisa_status.err:
            logger.error("Erro ao converter HTML para PDF com xhtml2pdf: %s", pisa_status.err)
            return None

        pdf_io.seek(0)
        return pdf_io.read()

    except ImportError as e:
        logger.error("Bibliotecas mammoth/xhtml2pdf não instaladas: %s", e)
    except Exception as exc:
        logger.error("Erro ao converter DOCX para PDF (fallback): %s", exc, exc_info=True)
        # continua para o fallback de reportlab

    # Fallback final: renderização simples com reportlab (texto/tabelas)
    try:
        from docx import Document
        from docx.table import Table as DocxTable
        from docx.text.paragraph import Paragraph as DocxParagraph
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        import html as html_lib

        def iter_block_items(parent):
            for child in parent.element.body.iterchildren():
                if child.tag.endswith('}p'):
                    yield DocxParagraph(child, parent)
                elif child.tag.endswith('}tbl'):
                    yield DocxTable(child, parent)

        def run_to_markup(run):
            text = html_lib.escape(run.text or '').replace('\n', '<br/>')
            if not text:
                return ''
            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"
            return text

        alignment_map = {
            None: TA_LEFT,
            0: TA_LEFT,
            1: TA_CENTER,
            2: TA_RIGHT,
            3: TA_JUSTIFY,
        }

        styles = getSampleStyleSheet()
        base_style = ParagraphStyle(
            'DocxBase',
            parent=styles['Normal'],
            fontName='Times-Roman',
            fontSize=12,
            leading=15,
            spaceAfter=6,
        )

        doc = Document(BytesIO(docx_bytes))
        buffer = BytesIO()
        pdf = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )
        flowables = []

        for block in iter_block_items(doc):
            if isinstance(block, DocxParagraph):
                paragraph_text = ''.join(run_to_markup(run) for run in block.runs).strip()
                if not paragraph_text:
                    flowables.append(Spacer(1, 8))
                    continue
                style = ParagraphStyle(
                    'DocxParagraph',
                    parent=base_style,
                    alignment=alignment_map.get(block.alignment, TA_LEFT),
                )
                flowables.append(Paragraph(paragraph_text, style))
            elif isinstance(block, DocxTable):
                table_data = []
                for row in block.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = html_lib.escape(cell.text or '').replace('\n', '<br/>')
                        row_cells.append(Paragraph(cell_text, base_style))
                    table_data.append(row_cells)
                if table_data:
                    table = Table(table_data, hAlign='LEFT')
                    table.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ]))
                    flowables.append(table)
                    flowables.append(Spacer(1, 8))

        if not flowables:
            return None

        pdf.build(flowables)
        buffer.seek(0)
        return buffer.read()
    except Exception as exc:
        logger.error("Erro ao converter DOCX para PDF (reportlab): %s", exc, exc_info=True)
        return None


def _build_cobranca_docx_bytes(
    processo,
    polo_passivo,
    contratos,
    processo_override=None,
    custas_override=None,
    custas_paragrafo=None,
    custas_parcelas=None,
    custas_valor_parcela=None
):
    contratos = sorted(contratos, key=lambda c: (c.numero_contrato or '', c.id))
    processo_override = processo_override or {}
    override_vara = processo_override.get('vara')
    override_uf = processo_override.get('uf')
    override_valor_causa = processo_override.get('valor_causa')
    dados = {
        'PARTE CONTRÁRIA': (polo_passivo.nome or '').upper(),
        'CPF': _format_cpf(polo_passivo.documento),
    }

    endereco_parts = parse_endereco(polo_passivo.endereco)
    for key in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']:
        dados[key] = _format_address_component(endereco_parts.get(key, '') or '')

    dados['E_FORO'] = (override_vara if override_vara is not None else processo.vara) or ''
    dados['H_FORO'] = (override_uf if override_uf is not None else processo.uf or '').upper()
    dados['UF'] = (override_uf if override_uf is not None else processo.uf or '').upper()
    dados['CONTRATO'] = _formatar_lista_contratos(contratos)

    total_valor = _get_total_contrato_value(contratos, processo, override_valor_causa)
    total_valor_causa = _get_total_valor_causa(contratos, processo, override_valor_causa)

    dados['VALOR'] = _format_currency_brl(total_valor)
    dados['VALOR POR EXTENSO'] = number_to_words_pt_br(total_valor)

    dados['VALOR DA CAUSA'] = _format_currency_brl(total_valor_causa)
    dados['VALOR DA CAUSA POR EXTENSO'] = number_to_words_pt_br(total_valor_causa)

    if custas_override is not None:
        custas_iniciais_25 = _to_decimal(custas_override).quantize(
            Decimal('0.01'),
            rounding=ROUND_HALF_UP
        )
        custas_iniciais_2 = custas_iniciais_25
    else:
        custas_iniciais_2 = (total_valor_causa * Decimal('0.02')).quantize(
            Decimal('0.01'),
            rounding=ROUND_HALF_UP
        )
        custas_iniciais_25 = (total_valor_causa * Decimal('0.025')).quantize(
            Decimal('0.01'),
            rounding=ROUND_HALF_UP
        )
    custas_extenso_2 = number_to_words_pt_br(custas_iniciais_2)
    custas_extenso_25 = number_to_words_pt_br(custas_iniciais_25)
    custas_formatadas_2 = _format_currency_brl(custas_iniciais_2)
    custas_formatadas_25 = _format_currency_brl(custas_iniciais_25)

    # Aliases genéricos seguem a regra corrente da peça (2,5%).
    custas_aliases = [
        'CUSTAS',
        'CUSTAS INICIAIS',
        'VALOR CUSTAS',
        'VALOR DAS CUSTAS',
    ]
    for alias in custas_aliases:
        dados[alias] = custas_formatadas_25

    custas_extenso_aliases = [
        'CUSTAS POR EXTENSO',
        'VALOR DAS CUSTAS POR EXTENSO',
    ]
    for alias in custas_extenso_aliases:
        dados[alias] = custas_extenso_25

    # Compatibilidade com templates legados (2%) e novos (2,5%).
    custas_2_aliases = [
        '2% DO VALOR DA CAUSA',
        '2 % DO VALOR DA CAUSA',
        '2,0% DO VALOR DA CAUSA',
        '2,0 % DO VALOR DA CAUSA',
        '2.0% DO VALOR DA CAUSA',
        '2.0 % DO VALOR DA CAUSA',
        'CUSTAS (2% DO VALOR DA CAUSA)',
        'CUSTAS (2 % DO VALOR DA CAUSA)',
    ]
    for alias in custas_2_aliases:
        dados[alias] = custas_formatadas_2

    custas_2_extenso_aliases = [
        '2% DO VALOR DA CAUSA POR EXTENSO',
        '2 % DO VALOR DA CAUSA POR EXTENSO',
        '2,0% DO VALOR DA CAUSA POR EXTENSO',
        '2,0 % DO VALOR DA CAUSA POR EXTENSO',
        '2.0% DO VALOR DA CAUSA POR EXTENSO',
        '2.0 % DO VALOR DA CAUSA POR EXTENSO',
    ]
    for alias in custas_2_extenso_aliases:
        dados[alias] = custas_extenso_2

    custas_25_aliases = [
        '2,5% DO VALOR DA CAUSA',
        '2,5 % DO VALOR DA CAUSA',
        '2.5% DO VALOR DA CAUSA',
        '2.5 % DO VALOR DA CAUSA',
        'CUSTAS (2,5% DO VALOR DA CAUSA)',
        'CUSTAS (2,5 % DO VALOR DA CAUSA)',
        'CUSTAS (2.5% DO VALOR DA CAUSA)',
        'CUSTAS (2.5 % DO VALOR DA CAUSA)',
    ]
    for alias in custas_25_aliases:
        dados[alias] = custas_formatadas_25

    custas_25_extenso_aliases = [
        '2,5% DO VALOR DA CAUSA POR EXTENSO',
        '2,5 % DO VALOR DA CAUSA POR EXTENSO',
        '2.5% DO VALOR DA CAUSA POR EXTENSO',
        '2.5 % DO VALOR DA CAUSA POR EXTENSO',
    ]
    for alias in custas_25_extenso_aliases:
        dados[alias] = custas_extenso_25

    parcelas_custas, valor_parcela = _resolve_custas_parcelamento(
        custas_iniciais_25,
        parcelas_override=custas_parcelas,
        valor_parcela_override=custas_valor_parcela
    )
    dados['X PARCELAS'] = str(parcelas_custas)
    dados['X PARCELAS POR EXTENSO'] = number_to_words_pt_br(
        parcelas_custas,
        feminine=True,
        include_currency=False,
        capitalize_first=False
    )

    dados['DATA DE HOJE'] = datetime.now().strftime("%d/%m/%Y")

    document = _load_template_document(DocumentoModelo.SlugChoices.COBRANCA_JUDICIAL, None)
    show_parcelamento = custas_iniciais_25 >= Decimal('1000')

    for section in document.sections:
        try:
            section.footer_distance = Cm(1.5)
        except Exception:
            pass

    if not show_parcelamento:
        for marker in (
            "DAS CUSTAS",
            "II. DAS CUSTAS",
            "Seja deferido o parcelamento das custas iniciais",
            "Considerando que a ora Exequente está assumindo a posição em milhares de",
            "O art. 98",
            "Para fins de transparência e controle",
            "A requerente requer o parcelamento em",
        ):
            _remove_paragraphs_containing(document, marker)
    elif custas_paragrafo:
        _replace_paragraphs_containing(
            document,
            "DAS CUSTAS",
            custas_paragrafo,
            first_only=True
        )
        _replace_paragraphs_containing(
            document,
            "II. DAS CUSTAS",
            custas_paragrafo,
            first_only=True
        )
        for marker in (
            "Seja deferido o parcelamento das custas iniciais",
            "Considerando que a ora Exequente está assumindo a posição em milhares de",
            "O art. 98",
            "Para fins de transparência e controle",
            "A requerente requer o parcelamento em",
        ):
            for paragraph in list(_iter_container_paragraphs(document)):
                text = paragraph.text or ''
                if marker in text and "DAS CUSTAS" not in text and "II. DAS CUSTAS" not in text:
                    _delete_paragraph(paragraph)

    _replace_placeholders_in_container(document, dados)
    for section in document.sections:
        _replace_placeholders_in_container(section.header, dados)
        _replace_placeholders_in_container(section.footer, dados)

    if show_parcelamento:
        _complete_cobranca_custas_paragraphs(
            document,
            custas_iniciais_25,
            parcelas_custas,
            valor_parcela
        )

    _apply_placeholder_styles(document)
    _bold_keywords_in_document(document, ['EXCELENTÍSSIMO(A)'])

    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream.getvalue()


def _parse_habilitacao_data(polo_passivo):
    endereco = polo_passivo.endereco or ''
    parts = parse_endereco(endereco)
    cidade = parts.get('F') or parts.get('E') or ''
    comarca = parts.get('F') or ''
    uf = parts.get('H') or ''
    return {
        'ENDERECO': endereco,
        'CIDADE': cidade,
        'COMARCA': comarca,
        'UF': uf
    }


def _prefer_non_empty_text(primary, fallback=''):
    if primary is None:
        return fallback
    if isinstance(primary, str):
        cleaned = primary.strip()
        if cleaned:
            return cleaned
        return fallback
    return primary


def _collect_missing_habilitacao_fields(processo, polo_passivo, processo_override=None):
    processo_override = processo_override or {}
    override_cnj = processo_override.get('cnj')
    override_vara = processo_override.get('vara')
    missing = []
    cnj_value = _prefer_non_empty_text(override_cnj, processo.cnj)
    if not (cnj_value and str(cnj_value).strip()):
        missing.append('Número de processo (CNJ)')
    vara_value = _prefer_non_empty_text(override_vara, processo.vara)
    if not (vara_value and str(vara_value).strip()):
        missing.append('Vara')
    endereco = polo_passivo.endereco or ''
    endereco_parts = parse_endereco(endereco)
    address_labels = {
        'A': 'Logradouro',
        'B': 'Número',
        'D': 'Bairro',
        'E': 'Cidade',
        'F': 'Comarca',
        'G': 'CEP',
        'H': 'UF'
    }
    endereco_missing = [
        address_labels[key]
        for key, label in address_labels.items()
        if not endereco_parts.get(key)
    ]
    if endereco_missing:
        missing.append('Endereço (' + ', '.join(endereco_missing) + ')')
    return missing


def _format_vara_text(vara_raw):
    if not vara_raw:
        return ''
    cleaned = re.sub(r'\s+', ' ', vara_raw).strip()
    digits = re.search(r'(\d+)', cleaned)
    if digits:
        num = digits.group(1)
        return f"{num}ª VARA"
    upper = cleaned.upper()
    return upper if 'VARA' in upper else f"{upper} VARA"


def _extract_comarca_from_vara(vara_raw):
    if not vara_raw:
        return ''
    cleaned = re.sub(r'\s+', ' ', str(vara_raw)).strip()
    match = re.search(r'\bde\s+(.+)$', cleaned, flags=re.IGNORECASE)
    if not match:
        return ''
    comarca = match.group(1).strip(" -")
    comarca = re.sub(r'\s*[-/]\s*[A-Z]{2}$', '', comarca, flags=re.IGNORECASE)
    return comarca.strip()


def _replace_with_style(document, pattern, replacement, uppercase=False, bold=False):
    if not replacement:
        return
    value = replacement.upper() if uppercase else replacement
    _replacePlaceholderStyled_(document, pattern, value, bold)


def _build_habilitacao_docx_bytes(processo, polo_passivo, processo_override=None):
    processo_override = processo_override or {}
    override_cnj = processo_override.get('cnj')
    override_vara = processo_override.get('vara')
    override_uf = processo_override.get('uf')
    replacements = _parse_habilitacao_data(polo_passivo)
    endereco = replacements.get('ENDERECO', '')
    cidade = replacements.get('CIDADE', '')
    comarca = replacements.get('COMARCA', '')
    uf = replacements.get('UF', '')
    vara_value = _prefer_non_empty_text(override_vara, processo.vara)
    comarca_header = _extract_comarca_from_vara(vara_value) or comarca or cidade
    cidade_foro = comarca_header or cidade
    uf_foro = _prefer_non_empty_text(override_uf, processo.uf) or uf or ''
    document = _load_template_document(DocumentoModelo.SlugChoices.HABILITACAO, None)

    _replace_with_style(
        document,
        '[VARA]',
        _format_vara_text(vara_value),
        uppercase=True
    )
    _replace_with_style(document, '[CIDADE]', cidade_foro, uppercase=True)
    _replace_with_style(document, '[UF]', uf_foro.upper(), uppercase=True)
    _replace_with_style(
        document,
        '[Processo]',
        _prefer_non_empty_text(override_cnj, processo.cnj) or '',
        uppercase=False
    )
    _replace_with_style(document, '[Polo Passivo]', polo_passivo.nome, uppercase=False)
    _replace_with_style(document, '[Polo Passivo MAIÚSCULAS]', polo_passivo.nome, uppercase=True)
    _replace_with_style(document, '[Polo Passivo TODAS MAIÚSCULAS E NEGRITO]', polo_passivo.nome, uppercase=True, bold=True)
    hoje = datetime.now()
    month_names = {
        '01': 'janeiro', '02': 'fevereiro', '03': 'março', '04': 'abril',
        '05': 'maio', '06': 'junho', '07': 'julho', '08': 'agosto',
        '09': 'setembro', '10': 'outubro', '11': 'novembro', '12': 'dezembro'
    }
    dia = hoje.strftime('%d')
    mes_extenso = month_names.get(hoje.strftime('%m'), '')
    ano = hoje.strftime('%Y')
    data_por_extenso = f"{dia} de {mes_extenso} de {ano}"
    _replace_with_style(document, '[DATA DE HOJE]', data_por_extenso, uppercase=False)

    saudacao = (
        f"EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA {_format_vara_text(vara_value)} "
        f"DA COMARCA DE {comarca_header.upper()} - {uf_foro.upper()}"
    )
    _replace_with_style(document, '[CABEÇALHO]', saudacao, uppercase=True, bold=True)
    _bold_paragraphs_containing(document, saudacao)

    local_date = f"{cidade_foro.capitalize()}/{uf_foro.upper()}, {data_por_extenso}"
    _replace_with_style(document, '[LOCAL_DATA]', local_date, uppercase=False)

    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream.getvalue()


def _only_digits(value):
    return re.sub(r'\D', '', str(value or ''))


def _fetch_extrato_titularidade(processo, polo_passivo, contratos, user):
    api_key = getattr(settings, 'JUDICIAL_API_KEY', None)
    if not api_key:
        return None

    cpf_digits = _only_digits(polo_passivo.documento)
    if len(cpf_digits) != 11:
        logger.warning("CPF inválido para extrato de titularidade: %s", polo_passivo.documento)
        return None

    contratos_numeros = [
        _only_digits(contrato.numero_contrato) for contrato in contratos
        if _only_digits(contrato.numero_contrato)
    ]
    if not contratos_numeros:
        return None

    include_param = quote(','.join(contratos_numeros))
    url = f'https://erp-api.nowlex.com/api/judicial/cpf/{cpf_digits}/pdf?include_contracts={include_param}'
    headers = {'X-API-Key': api_key}
    try:
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        contratos_label = _formatar_lista_contratos(contratos) or 'contratos'
        nome_parte = _extrair_primeiros_nomes(polo_passivo.nome or '', 2) or 'parte'
        file_name = f"05 - Extrato de Titularidade - {contratos_label} - {nome_parte}.pdf"

        arquivo_extrato = ProcessoArquivo(
            processo=processo,
            nome=file_name,
            enviado_por=user if user and user.is_authenticated else None,
        )
        arquivo_extrato.arquivo.save(_sanitize_filename(file_name), ContentFile(response.content), save=True)
        return arquivo_extrato.arquivo.url
    except requests.RequestException as exc:
        logger.warning("Erro ao buscar extrato de titularidade: %s", exc)
    except Exception as exc:
        logger.warning("Erro ao salvar extrato de titularidade: %s", exc, exc_info=True)
    return None

@login_required
@require_POST
def buscar_dados_escavador_view(request):
    """
    View que recebe uma requisição AJAX com um CNJ, busca os dados no Escavador
    e os retorna como JSON para preenchimento do formulário, SEM SALVAR.
    """
    cnj = request.POST.get('cnj')
    if not cnj:
        return JsonResponse({'status': 'error', 'message': 'CNJ não fornecido.'}, status=400)

    try:
        # 1. Buscar os dados brutos da API
        dados_api = buscar_processo_por_cnj(cnj)
        if not dados_api:
            return JsonResponse({
                'status': 'error',
                'message': f'Não foi possível encontrar dados para o CNJ {cnj}.'
            }, status=404)

        # --- Início do Bloco de Processamento Seguro ---
        
        # 2. Preparar os dados para o formulário
        fontes_list = dados_api.get('fontes', [])
        fonte_principal = fontes_list[0] if fontes_list else {}

        # Trata o valor da causa
        valor_causa = Decimal('0.00')
        if fonte_principal:
            valor_causa_raw = fonte_principal.get('capa', {}).get('valor_causa', {}).get('valor_formatado')
            if valor_causa_raw is not None:
                valor_causa_str = str(valor_causa_raw).replace('R$', '').replace('.', '').replace(',', '.').strip()
                if valor_causa_str:
                    valor_causa = Decimal(valor_causa_str)

        # Trata o status (classe processual)
        status_id = None
        status_nome = None
        if fonte_principal:
            nome_classe_processual = fonte_principal.get('capa', {}).get('classe')
            if nome_classe_processual:
                normalized_name = re.sub(r'\s*\(\d+\)$', '', nome_classe_processual).strip()

                status, created = StatusProcessual.get_or_create_normalized(
                    normalized_name,
                    defaults={'ordem': 0}
                )
                status_id = status.id
                status_nome = status.nome

        # Prepara a lista de partes
        partes_para_formulario = collect_partes_from_escavador_payload(dados_api)

        # Prepara a lista de andamentos
        andamentos_para_formulario = []
        if fonte_principal:
            for andamento_api in dados_api.get('movimentacoes', []):
                andamentos_para_formulario.append({
                    'data': andamento_api.get('data'),
                    'descricao': andamento_api.get('conteudo'),
                })

        # 3. Montar o dicionário de resposta final
        dados_completos = {
            'status': 'success',
            'message': 'Dados encontrados! Revise e salve o formulário.',
            'processo': {
                'uf': dados_api.get('estado_origem', {}).get('sigla', ''),
                'vara': fonte_principal.get('capa', {}).get('orgao_julgador', ''),
                'tribunal': fonte_principal.get('tribunal', {}).get('nome', ''),
                'valor_causa': f'{valor_causa:.2f}',
                'status_id': status_id,
                'status_nome': status_nome,
            },
            'partes': partes_para_formulario,
            'andamentos': andamentos_para_formulario,
        }
        return JsonResponse(dados_completos)

    except Exception as e:
        # Captura QUALQUER erro durante a busca ou processamento dos dados
        logger.error(f"Erro ao processar CNJ {cnj}: {e}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': f'Ocorreu um erro interno ao processar os dados da API: {e}'
        }, status=500)


@staff_member_required
@require_POST
@transaction.atomic
def merge_status_view(request):
    """
    View para mesclar dois Status Processuais.
    Atualiza todos os Processos Judiciais do status de origem para o de destino.
    """
    try:
        source_id = int(request.POST.get('source_id'))
        target_id = int(request.POST.get('target_id'))

        if source_id == target_id:
            return JsonResponse({'status': 'error', 'message': 'Os status de origem e destino não podem ser os mesmos.'}, status=400)

        source_status = get_object_or_404(StatusProcessual, pk=source_id)
        target_status = get_object_or_404(StatusProcessual, pk=target_id)

        affected_processes_count = ProcessoJudicial.objects.filter(status=source_status).count()
        ProcessoJudicial.objects.filter(status=source_status).update(status=target_status)
        
        source_status.ativo = False
        source_status.save()
        
        message = f'{affected_processes_count} processo(s) foram atualizados. O status "{source_status.nome}" foi mesclado e inativado.'
        return JsonResponse({'status': 'success', 'message': message})

    except (KeyError, ValueError, TypeError):
        return JsonResponse({'status': 'error', 'message': 'Dados inválidos fornecidos.'}, status=400)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Ocorreu um erro inesperado: {e}'}, status=500)


@login_required
def lista_processos(request):
    """
    Busca todos os processos no banco e os envia para o template de lista.
    """
    processos = filter_processos_queryset_for_user(
        ProcessoJudicial.objects.all(),
        request.user,
    ).order_by('-id')
    contexto = {
        'processos': processos
    }
    return render(request, 'contratos/lista_processos.html', contexto)


@login_required
def detalhe_processo(request, pk):
    """
    Busca um processo específico pelo seu ID (pk) e envia para o template de detalhe.
    """
    processo = get_object_or_404(
        filter_processos_queryset_for_user(ProcessoJudicial.objects.all(), request.user),
        pk=pk,
    )
    contratos_do_processo = processo.contratos.all()
    
    contexto = {
        'processo': processo,
        'contratos': contratos_do_processo,
    }
    return render(request, 'contratos/detalhe_processo.html', contexto)

@login_required
@require_GET
def get_analysis_types(request):
    try:
        tipos = TipoAnaliseObjetiva.objects.filter(ativo=True).order_by('nome')
    except (ProgrammingError, OperationalError):
        return JsonResponse({'status': 'success', 'types': []})
    return JsonResponse({
        'status': 'success',
        'types': [
            {
                'id': tipo.id,
                'nome': tipo.nome,
                'slug': tipo.slug,
                'hashtag': tipo.hashtag,
                'versao': tipo.versao,
            }
            for tipo in tipos
        ]
    })


@login_required
@require_GET
def get_decision_tree_data(request):
    """
    Retorna a estrutura da árvore de decisão (questões e opções) como JSON,
    mesclando a configuração nativa com as do banco de dados.
    """
    tipo_id = request.GET.get('tipo_id')
    tipo_slug = request.GET.get('tipo_slug')
    tipo = None
    if tipo_id:
        try:
            tipo = TipoAnaliseObjetiva.objects.get(pk=int(tipo_id))
        except (ProgrammingError, OperationalError):
            return JsonResponse({'status': 'error', 'message': 'Tipos de análise ainda não disponíveis. Rode as migrações.'}, status=400)
        except (TipoAnaliseObjetiva.DoesNotExist, ValueError, TypeError):
            return JsonResponse({'status': 'error', 'message': 'Tipo de análise inválido.'}, status=400)
    elif tipo_slug:
        try:
            tipo = TipoAnaliseObjetiva.objects.filter(slug=tipo_slug).first()
        except (ProgrammingError, OperationalError):
            return JsonResponse({'status': 'error', 'message': 'Tipos de análise ainda não disponíveis. Rode as migrações.'}, status=400)

    if not tipo:
        try:
            tipo = (
                TipoAnaliseObjetiva.objects.filter(slug='novas-monitorias').first()
                or TipoAnaliseObjetiva.objects.filter(ativo=True).order_by('nome').first()
            )
        except (ProgrammingError, OperationalError):
            return JsonResponse({'status': 'error', 'message': 'Tipos de análise ainda não disponíveis. Rode as migrações.'}, status=400)

    if not tipo:
        return JsonResponse({'status': 'error', 'message': 'Nenhum tipo de análise configurado.'}, status=404)

    # 1. Inicia com a configuração nativa
    # Cria uma cópia profunda para evitar modificar o dicionário original
    current_tree_config = copy.deepcopy(DECISION_TREE_CONFIG) if tipo.slug == 'novas-monitorias' else {}
    
    primeira_questao_chave_db = None

    # 2. Mescla/Sobrescreve com as questões configuradas no banco de dados
    db_questoes = (
        QuestaoAnalise.objects
        .filter(tipo_analise=tipo, ativo=True)
        .prefetch_related('opcoes')
    )
    for questao_db in db_questoes:
        if not questao_db.chave: # Chave é essencial para mesclar
            continue

        q_data = {
            'id': questao_db.id,
            'texto_pergunta': questao_db.texto_pergunta,
            'chave': questao_db.chave,
            'tipo_campo': questao_db.tipo_campo,
            'is_primeira_questao': questao_db.is_primeira_questao,
            'habilita_supervisao': bool(questao_db.habilita_supervisao),
            'ordem': questao_db.ordem,
            'opcoes': []
        }

        for opcao_db in questao_db.opcoes.filter(ativo=True).all().order_by('id'):
            prox = opcao_db.proxima_questao
            prox_same_tipo = bool(
                prox
                and getattr(prox, 'tipo_analise_id', None)
                and getattr(questao_db, 'tipo_analise_id', None)
                and prox.tipo_analise_id == questao_db.tipo_analise_id
            )
            o_data = {
                'id': opcao_db.id,
                'texto_resposta': opcao_db.texto_resposta,
                'proxima_questao_id': prox.id if (prox and prox_same_tipo) else None,
                'proxima_questao_chave': prox.chave if (prox and prox_same_tipo and prox.chave) else None,
                'proxima_questao_texto': prox.texto_pergunta if (prox and prox_same_tipo) else None,
            }
            q_data['opcoes'].append(o_data)

        # Para campos que não são dropdown, aceita um "fluxo linear" usando a primeira opção
        # que apontar uma próxima questão.
        if questao_db.tipo_campo != 'OPCOES':
            next_opt = next(
                (
                    o
                    for o in q_data.get('opcoes', [])
                    if o and o.get('proxima_questao_chave')
                ),
                None,
            )
            if next_opt:
                q_data['proxima_questao_chave'] = next_opt.get('proxima_questao_chave')
                q_data['proxima_questao_texto'] = next_opt.get('proxima_questao_texto')
        
        # Se a questão já existe na config nativa, sobrescreve seus campos
        # e mescla as opções. Caso contrário, adiciona a questão.
        if questao_db.chave in current_tree_config:
            # Sobrescreve atributos da questão
            current_tree_config[questao_db.chave].update({
                'id': q_data['id'],
                'texto_pergunta': q_data['texto_pergunta'],
                'tipo_campo': q_data['tipo_campo'],
                'is_primeira_questao': q_data['is_primeira_questao'],
                'habilita_supervisao': q_data['habilita_supervisao'],
                'ordem': q_data['ordem'],
            })
            if q_data.get('proxima_questao_chave'):
                current_tree_config[questao_db.chave]['proxima_questao_chave'] = q_data.get('proxima_questao_chave')
                if q_data.get('proxima_questao_texto'):
                    current_tree_config[questao_db.chave]['proxima_questao_texto'] = q_data.get('proxima_questao_texto')
            # Substitui as opções pelas do banco de dados
            current_tree_config[questao_db.chave]['opcoes'] = q_data['opcoes']
        else:
            current_tree_config[questao_db.chave] = q_data
        
        if questao_db.is_primeira_questao:
            primeira_questao_chave_db = questao_db.chave
    
    # 3. Determina a chave da primeira questão
    # Prioriza a primeira questão definida no banco de dados
    if primeira_questao_chave_db:
        final_primeira_questao_chave = primeira_questao_chave_db
    else:
        # Se não houver no banco, usa a da configuração nativa
        # Assume que há apenas uma 'is_primeira_questao = True' na config nativa
        for chave, questao_data in current_tree_config.items():
            if questao_data.get("is_primeira_questao"):
                final_primeira_questao_chave = chave
                break
        else: # Se não encontrar nem na nativa
            return JsonResponse({'status': 'error', 'message': 'Nenhuma questão inicial configurada no banco de dados ou na configuração nativa.'}, status=404)

    return JsonResponse({
        'status': 'success',
        'analysis_type': {
            'id': tipo.id,
            'nome': tipo.nome,
            'slug': tipo.slug,
            'hashtag': tipo.hashtag,
            'versao': tipo.versao,
        },
        'primeira_questao_chave': final_primeira_questao_chave,
        'tree_data': current_tree_config
    })

@login_required
@require_GET
def get_processo_contratos_api(request, processo_id):
    """
    Retorna uma lista de contratos associados a um Processo Judicial específico como JSON.
    """
    try:
        processo = get_object_or_404(ProcessoJudicial, pk=processo_id)
        contratos = processo.contratos.all().order_by('numero_contrato')
        
        contratos_data = [
            {
                'id': contrato.id,
                'numero_contrato': contrato.numero_contrato or f'Contrato sem número ({contrato.id})',
                'valor_total_devido': str(contrato.valor_total_devido) if contrato.valor_total_devido else None,
            }
            for contrato in contratos
        ]
        
        return JsonResponse({'status': 'success', 'contratos': contratos_data})
    except ProcessoJudicial.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Processo Judicial não encontrado.'}, status=404)
    except Exception as e:
        logger.error(f"Erro ao buscar contratos para processo_id {processo_id}: {e}", exc_info=True)
        return JsonResponse({'status': 'error', 'message': f'Ocorreu um erro interno: {e}'}, status=500)
# ==============================================================================
# FUNÇÕES E VIEWS PARA GERAÇÃO DA PEÇA MONITÓRIA
# ==============================================================================
from docx import Document
from io import BytesIO
from datetime import datetime
import os
from django.conf import settings

# Helper para parsear o endereço
def parse_endereco(endereco_str):
    parts = {
        'A': '', 'B': '', 'C': '', 'D': '', 'E': '',
        'F': '', 'G': '', 'H': ''
    }
    if not endereco_str:
        return parts

    # Expressão regular para capturar CHAVE: VALOR
    # Garante que o valor pode conter hífens e espaços
    matches = re.findall(r'([A-H]):\s*([^:]+?)(?=\s*-\s*[A-H]:|\s*$)', endereco_str)
    
    for key, value in matches:
        value = value.strip()
        # Limpa 'None' e 'null' da string, que podem vir de campos vazios
        if value.lower() == 'none' or value.lower() == 'null':
            value = ''
        parts[key.strip()] = value
    return parts


def _resolve_peticao_cnj_entry(processo, entry_id=None, cnj_value=None):
    if not processo:
        return None
    if entry_id not in (None, ''):
        try:
            entry_id_int = int(entry_id)
        except (TypeError, ValueError):
            entry_id_int = None
        if entry_id_int:
            entry = processo.numeros_cnj.filter(pk=entry_id_int).first()
            if entry:
                return entry
    if cnj_value:
        digits = re.sub(r'\D', '', str(cnj_value or ''))
        if digits:
            for entry in processo.numeros_cnj.all():
                if re.sub(r'\D', '', entry.cnj or '') == digits:
                    return entry
    return None


@login_required
@require_POST
def update_peticao_dados(request, processo_id=None):
    processo_id = processo_id or request.POST.get('processo_id')
    try:
        processo_id_int = int(processo_id)
    except (TypeError, ValueError):
        return JsonResponse({'error': 'ID do processo inválido.'}, status=400)

    processo = get_object_or_404(ProcessoJudicial, pk=processo_id_int)
    try:
        payload = json.loads(request.body.decode('utf-8') or '{}')
    except Exception:
        payload = {}

    source = str(payload.get('source') or 'base').strip().lower()
    cnj_entry_id = payload.get('cnj_entry_id')
    cnj_value = (payload.get('cnj') or '').strip()
    uf_value = (payload.get('uf') or '').strip().upper()[:2]
    vara_value = (payload.get('vara') or '').strip()
    tribunal_value = (payload.get('tribunal') or '').strip()
    valor_causa_value = _parse_decimal_input(payload.get('valor_causa'))
    endereco_parts = payload.get('endereco_parts')
    endereco_raw = payload.get('endereco_raw')
    update_base_cnj = bool(payload.get('update_base_cnj'))

    updated_entry = None
    updated_processo = None

    if source == 'cnj':
        entry = _resolve_peticao_cnj_entry(processo, entry_id=cnj_entry_id, cnj_value=cnj_value)
        if not entry:
            return JsonResponse({'error': 'Número CNJ não encontrado no cadastro.'}, status=404)
        if cnj_value:
            entry.cnj = cnj_value
        entry.uf = uf_value
        entry.vara = vara_value
        entry.tribunal = tribunal_value
        entry.valor_causa = valor_causa_value
        entry.save()
        updated_entry = {
            'id': entry.pk,
            'cnj': entry.cnj or '',
            'uf': entry.uf or '',
            'vara': entry.vara or '',
            'tribunal': entry.tribunal or '',
            'valor_causa': str(entry.valor_causa or '') if entry.valor_causa is not None else '',
        }
    else:
        processo.uf = uf_value
        processo.vara = vara_value
        processo.tribunal = tribunal_value
        processo.valor_causa = valor_causa_value
        if update_base_cnj and cnj_value:
            processo.cnj = cnj_value
        processo.save()
        updated_processo = {
            'cnj': processo.cnj or '',
            'uf': processo.uf or '',
            'vara': processo.vara or '',
            'tribunal': processo.tribunal or '',
            'valor_causa': str(processo.valor_causa or '') if processo.valor_causa is not None else '',
        }

    endereco_text = None
    if isinstance(endereco_parts, dict):
        endereco_text = _build_endereco_from_parts(endereco_parts)
    elif isinstance(endereco_raw, str):
        endereco_text = endereco_raw.strip()

    updated_parte_id = None
    if endereco_text is not None:
        polo_passivo = None
        if payload.get('polo_passivo_id'):
            try:
                polo_passivo = processo.partes_processuais.filter(
                    pk=int(payload.get('polo_passivo_id'))
                ).first()
            except (TypeError, ValueError):
                polo_passivo = None
        if polo_passivo is None:
            polo_passivo = processo.partes_processuais.filter(tipo_polo='PASSIVO').first()
        if polo_passivo:
            polo_passivo.endereco = endereco_text
            polo_passivo.save(update_fields=['endereco'])
            updated_parte_id = polo_passivo.pk

    return JsonResponse({
        'status': 'success',
        'source': source,
        'processo': updated_processo,
        'cnj_entry': updated_entry,
        'endereco': endereco_text,
        'polo_passivo_id': updated_parte_id,
    })


ADDRESS_LOWERCASE_WORDS = {
    'a', 'à', 'ao', 'aos', 'as',
    'da', 'das', 'de', 'do', 'dos',
    'em', 'no', 'na', 'nos', 'nas',
    'para', 'por', 'pelo', 'pela', 'pelos', 'pelas',
    'e', 'com', 'sem'
}
_ADDRESS_TOKEN_PATTERN = re.compile(r'^([^A-Za-zÀ-ÿ]*)([A-Za-zÀ-ÿ]+)([^A-Za-zÀ-ÿ]*)$', re.UNICODE)


def _format_address_component(value):
    if not value:
        return ''
    tokens = value.split()
    formatted = []
    for token in tokens:
        match = _ADDRESS_TOKEN_PATTERN.match(token)
        if not match:
            formatted.append(token)
            continue
        prefix, core, suffix = match.groups()
        if len(core) <= 1:
            formatted.append(token)
            continue
        lower_core = core.lower()
        if lower_core in ADDRESS_LOWERCASE_WORDS:
            normalized = lower_core
        elif core.isupper() and len(core) <= 3:
            normalized = core
        else:
            normalized = lower_core.capitalize()
        formatted.append(f"{prefix}{normalized}{suffix}")
    return ' '.join(formatted)

# Helper para converter número para extenso (simplificado)
# Para uma solução robusta, usar uma biblioteca ou implementar mais completo.
def number_to_words_pt_br(num, feminine=False, include_currency=True, capitalize_first=True):
    try:
        num_decimal = num if isinstance(num, Decimal) else Decimal(str(num))
    except (InvalidOperation, ValueError, TypeError):
        return str(num)

    unidades_masc = ['', 'um', 'dois', 'três', 'quatro', 'cinco', 'seis', 'sete', 'oito', 'nove']
    unidades_fem = ['', 'uma', 'duas', 'três', 'quatro', 'cinco', 'seis', 'sete', 'oito', 'nove']
    dezena = ['', 'dez', 'vinte', 'trinta', 'quarenta', 'cinquenta', 'sessenta', 'setenta', 'oitenta', 'noventa']
    dez_a_dezenove = ['dez', 'onze', 'doze', 'treze', 'catorze', 'quinze', 'dezesseis', 'dezessete', 'dezoito', 'dezenove']
    centenas = ['', 'cento', 'duzentos', 'trezentos', 'quatrocentos', 'quinhentos', 'seiscentos', 'setecentos', 'oitocentos', 'novecentos']
    unidades = unidades_fem if feminine else unidades_masc

    def _num_to_words_chunk(n):
        s = ''
        n = int(n)

        if n >= 100:
            if n == 100:
                s += 'cem'
            else:
                s += centenas[n // 100]
            n %= 100
            if n > 0:
                s += ' e '

        if 10 <= n < 20:
            s += dez_a_dezenove[n - 10]
            return s

        if n >= 20:
            s += dezena[n // 10]
            n %= 10
            if n > 0:
                s += ' e '

        if n > 0:
            s += unidades[n]

        return s

    def _process_triplet(triplet, scale):
        triplet = int(triplet)
        if triplet == 0:
            return ''
        words = _num_to_words_chunk(triplet)

        if scale == 1:  # Mil
            return 'mil' if words == 'um' else f"{words} mil"
        if scale == 2:  # Milhão
            return 'um milhão' if words == 'um' else f"{words} milhões"
        return words

    inteiro = int(num_decimal)
    words_list = []
    if inteiro == 0:
        words_list.append('zero')
    else:
        temp_int = inteiro
        idx = 0
        while temp_int > 0:
            chunk = temp_int % 1000
            if chunk > 0:
                words = _process_triplet(chunk, idx)
                if words:
                    words_list.insert(0, words)
            temp_int //= 1000
            idx += 1

    inteiro_words = ' '.join(filter(None, words_list)) or 'zero'
    inteiro_phrase = f"{inteiro_words} reais" if include_currency else inteiro_words

    decimal_part = int(((num_decimal - inteiro) * 100).quantize(Decimal('1'), rounding=ROUND_HALF_UP))
    if decimal_part > 0:
        centavos_text = _num_to_words_chunk(decimal_part).strip()
        centavos_phrase = f"{centavos_text} centavos" if centavos_text else 'centavos'
        inteiro_phrase = f"{inteiro_phrase} e {centavos_phrase}"

    return inteiro_phrase.capitalize() if capitalize_first else inteiro_phrase

@login_required
@require_POST
def generate_monitoria_petition(request, processo_id=None):
    # aceita tanto o ID vindo na URL quanto no POST (fallback)
    processo_id = processo_id or request.POST.get('processo_id')
    
    try:
        processo_id_int = int(processo_id)
    except (TypeError, ValueError):
        return HttpResponse("ID do processo inválido.", status=400)

    try:
        processo = get_object_or_404(ProcessoJudicial, pk=processo_id_int)
        analise = processo.analise_processo  # OneToOneField
    except ProcessoJudicial.DoesNotExist:
        return HttpResponse("Processo Judicial não encontrado.", status=404)
    except Exception as e:
        logger.error(f"Erro ao buscar análise do processo {processo_id_int}: {e}", exc_info=True)
        return HttpResponse(f"Erro ao buscar dados do processo: {e}", status=500)

    # Buscar a parte passiva
    polo_passivo = processo.partes_processuais.filter(tipo_polo='PASSIVO').first()
    if not polo_passivo:
        return HttpResponse("Polo passivo não encontrado para este processo.", status=404)
    
    # Buscar contratos selecionados para monitória (POST tem prioridade)
    contratos_para_monitoria_ids = []
    try:
        posted_json = request.POST.get('contratos_para_monitoria')
        if posted_json:
            contratos_para_monitoria_ids = json.loads(posted_json)
    except (TypeError, json.JSONDecodeError):
        contratos_para_monitoria_ids = []

    if not contratos_para_monitoria_ids:
        contratos_para_monitoria_ids = analise.respostas.get('contratos_para_monitoria', [])

    contratos_para_monitoria_ids = _parse_contract_ids(contratos_para_monitoria_ids)

    contratos_monitoria_qs = Contrato.objects.filter(id__in=contratos_para_monitoria_ids).select_related('processo')
    contratos_monitoria = list(contratos_monitoria_qs)

    if not contratos_monitoria:
        return HttpResponse("Nenhum contrato selecionado para monitória na análise deste processo.", status=404)

    try:
        cnj_entry = _resolve_peticao_cnj_entry(
            processo,
            entry_id=request.POST.get('peticao_cnj_entry_id'),
        )
        custas_override = _parse_decimal_input(request.POST.get('custas_total'))
        custas_paragrafo = _sanitize_monitoria_custas_paragrafo(
            request.POST.get('custas_paragrafo') or ''
        )
        custas_parcelas = _parse_int_input(request.POST.get('custas_parcelas'))
        custas_valor_parcela = _parse_decimal_input(request.POST.get('custas_valor_parcela'))
        processo_override = {}
        if cnj_entry:
            processo_override = {
                'cnj': cnj_entry.cnj or '',
                'uf': cnj_entry.uf or '',
                'vara': cnj_entry.vara or '',
                'tribunal': cnj_entry.tribunal or '',
                'valor_causa': cnj_entry.valor_causa,
            }

        docx_bytes = _build_docx_bytes_common(
            processo,
            polo_passivo,
            contratos_monitoria,
            processo_override=processo_override,
            custas_override=custas_override,
            custas_paragrafo=custas_paragrafo,
            custas_parcelas=custas_parcelas,
            custas_valor_parcela=custas_valor_parcela,
        )
        base_filename = _build_monitoria_base_filename(polo_passivo, contratos_monitoria)

        monitoria_info = {}
        docx_url = ''
        arquivo_pdf = None
        arquivo_docx = None

        try:
            docx_name = f"{base_filename}.docx"
            docx_file = ContentFile(docx_bytes)
            arquivo_docx = ProcessoArquivo(
                processo=processo,
                nome=docx_name,
                enviado_por=request.user if request.user.is_authenticated else None,
            )
            arquivo_docx.arquivo.save(docx_name, docx_file, save=True)
            docx_url = arquivo_docx.arquivo.url
        except Exception as exc:
            logger.error("Erro ao salvar DOCX da monitória: %s", exc, exc_info=True)
            docx_url = ''

        if not docx_url:
            return HttpResponse("Falha ao salvar o DOCX/PDF gerado nos Arquivos.", status=500)

        monitoria_info = {
            "ok": False,
            "pdf_url": '',
            "pdf_pending": True,
            "docx_download_url": docx_url or request.build_absolute_uri(
                reverse('contratos:generate_monitoria_docx', kwargs={'processo_id': processo_id_int})
            ),
        }

        extrato_result = generate_extrato_titularidade(
            processo=processo,
            cpf_value=polo_passivo.documento,
            contratos=contratos_monitoria,
            parte_name=polo_passivo.nome,
            usuario=request.user if request.user.is_authenticated else None
        )
        # Usa o arquivo PDF se existir, senão usa o DOCX para obter o dest_path
        if arquivo_docx:
            dest_path = os.path.dirname(arquivo_docx.arquivo.name or '')
        else:
            dest_path = ''

        response_payload = {
            "status": "success",
            "message": "Petição gerada - Salva em Arquivos.",
            "monitoria": monitoria_info,
            "extrato": extrato_result,
            "dest_path": dest_path,
        }
        return JsonResponse(response_payload)

    except Exception as e:
        logger.error(f"Erro ao gerar a petição para o processo {processo_id}: {e}", exc_info=True)
        return HttpResponse(f"Erro ao gerar a petição: {e}", status=500)


@login_required
@require_POST
def generate_cobranca_judicial_petition(request, processo_id=None):
    processo_id = processo_id or request.POST.get('processo_id')
    
    try:
        processo_id_int = int(processo_id)
    except (TypeError, ValueError):
        return HttpResponse("ID do processo inválido.", status=400)

    try:
        processo = get_object_or_404(ProcessoJudicial, pk=processo_id_int)
        analise = processo.analise_processo
    except ProcessoJudicial.DoesNotExist:
        return HttpResponse("Processo Judicial não encontrado.", status=404)
    except Exception as e:
        logger.error(f"Erro ao buscar análise do processo {processo_id_int}: {e}", exc_info=True)
        return HttpResponse(f"Erro ao buscar dados do processo: {e}", status=500)

    polo_passivo = processo.partes_processuais.filter(tipo_polo='PASSIVO').first()
    if not polo_passivo:
        return HttpResponse("Polo passivo não encontrado para este processo.", status=404)

    contratos_para_cobranca_ids = []
    try:
        posted_json = request.POST.get('contratos_para_monitoria')
        if posted_json:
            contratos_para_cobranca_ids = json.loads(posted_json)
    except (TypeError, json.JSONDecodeError):
        contratos_para_cobranca_ids = []

    contratos_queryset = processo.contratos.all()
    if contratos_para_cobranca_ids:
        contratos_queryset = contratos_queryset.filter(id__in=contratos_para_cobranca_ids)
    else:
        fallback_ids = getattr(analise, 'respostas', {}).get('contratos_para_monitoria', [])
        if fallback_ids:
            contratos_queryset = contratos_queryset.filter(id__in=fallback_ids)

    contratos_lista = list(contratos_queryset)
    if not contratos_lista:
        return HttpResponse("Nenhum contrato disponível para gerar a cobrança judicial.", status=404)

    if not polo_passivo.endereco:
        return HttpResponse("Endereço da parte passiva não informado.", status=400)

    extrato_result = None
    try:
        cnj_entry = _resolve_peticao_cnj_entry(
            processo,
            entry_id=request.POST.get('peticao_cnj_entry_id'),
        )
        custas_override = _parse_decimal_input(request.POST.get('custas_total'))
        custas_paragrafo = request.POST.get('custas_paragrafo') or ''
        custas_parcelas = _parse_int_input(request.POST.get('custas_parcelas'))
        custas_valor_parcela = _parse_decimal_input(request.POST.get('custas_valor_parcela'))
        processo_override = {}
        if cnj_entry:
            processo_override = {
                'cnj': cnj_entry.cnj or '',
                'uf': cnj_entry.uf or '',
                'vara': cnj_entry.vara or '',
                'tribunal': cnj_entry.tribunal or '',
                'valor_causa': cnj_entry.valor_causa,
            }

        docx_bytes = _build_cobranca_docx_bytes(
            processo,
            polo_passivo,
            contratos_lista,
            processo_override=processo_override,
            custas_override=custas_override,
            custas_paragrafo=custas_paragrafo,
            custas_parcelas=custas_parcelas,
            custas_valor_parcela=custas_valor_parcela,
        )
        base_filename = _build_cobranca_base_filename(polo_passivo, contratos_lista)
        docx_url = ''

        docx_saved = False
        try:
            docx_name = f"{base_filename}.docx"
            docx_file = ContentFile(docx_bytes)
            arquivo_docx = ProcessoArquivo(
                processo=processo,
                nome=docx_name,
                enviado_por=request.user if request.user.is_authenticated else None,
            )
            arquivo_docx.arquivo.save(docx_name, docx_file, save=True)
            docx_url = arquivo_docx.arquivo.url
            docx_saved = True
        except Exception as exc:
            logger.error("Erro ao salvar DOCX da cobrança judicial: %s", exc, exc_info=True)

        if not docx_saved:
            return HttpResponse("Falha ao salvar o DOCX/PDF gerado nos Arquivos.", status=500)
    except FileNotFoundError as fe:
        logger.error("Template de cobrança não encontrado: %s", fe)
        return HttpResponse(str(fe), status=500)
    except Exception as exc:
        logger.error(f"Erro ao gerar petição de cobrança para o processo {processo_id}: {exc}", exc_info=True)
        return HttpResponse(f"Erro ao gerar a petição de cobrança: {exc}", status=500)

    extrato_result = generate_extrato_titularidade(
        processo=processo,
        cpf_value=polo_passivo.documento,
        contratos=contratos_lista,
        parte_name=polo_passivo.nome,
        usuario=request.user if request.user.is_authenticated else None
    )

    cobranca_info = {
        "ok": False,
        "pdf_url": '',
        "pdf_pending": True,
        "docx_url": docx_url,
    }

    response_payload = {
        "status": "success",
        "message": "Petição de cobrança gerada - Salva em Arquivos.",
        "cobranca": cobranca_info,
        "extrato": extrato_result,
    }
    if extrato_result and isinstance(extrato_result, dict):
        pdf_url_extrato = extrato_result.get("pdf_url")
        if pdf_url_extrato:
            response_payload["extrato_url"] = pdf_url_extrato
        response_payload["message"] += " Extrato de titularidade salvo."

    return JsonResponse(response_payload)


@login_required
@require_POST
def generate_habilitacao_petition(request, processo_id=None):
    processo_id = processo_id or request.POST.get('processo_id')

    try:
        processo_id_int = int(processo_id)
    except (TypeError, ValueError):
        return HttpResponse("ID do processo inválido.", status=400)

    try:
        processo = get_object_or_404(ProcessoJudicial, pk=processo_id_int)
    except ProcessoJudicial.DoesNotExist:
        return HttpResponse("Processo Judicial não encontrado.", status=404)
    except Exception as exc:
        logger.error(f"Erro ao buscar o processo {processo_id_int}: {exc}", exc_info=True)
        return HttpResponse(f"Erro ao buscar o processo: {exc}", status=500)

    polo_passivo = processo.partes_processuais.filter(tipo_polo='PASSIVO').first()
    if not polo_passivo:
        return HttpResponse("Polo passivo não encontrado para este processo.", status=404)

    cnj_entry = _resolve_peticao_cnj_entry(
        processo,
        entry_id=request.POST.get('peticao_cnj_entry_id'),
    )
    processo_override = {}
    if cnj_entry:
        processo_override = {
            'cnj': cnj_entry.cnj or '',
            'uf': cnj_entry.uf or '',
            'vara': cnj_entry.vara or '',
            'tribunal': cnj_entry.tribunal or '',
            'valor_causa': cnj_entry.valor_causa,
        }

    missing_fields = _collect_missing_habilitacao_fields(processo, polo_passivo, processo_override)
    if missing_fields:
        message = (
            'Não foi possível gerar a habilitação porque faltam os seguintes dados no cadastro (aba Partes): '
            + '; '.join(missing_fields)
            + '.'
        )
        return JsonResponse({'message': message}, status=422)

    try:
        docx_bytes = _build_habilitacao_docx_bytes(processo, polo_passivo, processo_override)
        base_filename = _build_habilitacao_base_filename(
            polo_passivo,
            processo,
            cnj_reference=processo_override.get('cnj') if processo_override else None,
        )
        docx_url = ''

        docx_saved = False
        try:
            docx_name = f"{base_filename}.docx"
            docx_file = ContentFile(docx_bytes)
            arquivo_docx = ProcessoArquivo(
                processo=processo,
                nome=docx_name,
                enviado_por=request.user if request.user.is_authenticated else None,
            )
            arquivo_docx.arquivo.save(docx_name, docx_file, save=True)
            docx_url = arquivo_docx.arquivo.url
            docx_saved = True
        except Exception as exc:
            logger.error("Erro ao salvar DOCX da habilitação: %s", exc, exc_info=True)

        if not docx_saved:
            return HttpResponse("Falha ao salvar o DOCX/PDF gerado nos Arquivos.", status=500)
    except FileNotFoundError as fe:
        logger.error("Template de habilitação não encontrado: %s", fe)
        return HttpResponse(str(fe), status=500)
    except Exception as exc:
        logger.error(f"Erro ao gerar petição de habilitação para o processo {processo_id}: {exc}", exc_info=True)
        return HttpResponse(f"Erro ao gerar a petição de habilitação: {exc}", status=500)

    habilitacao_info = {
        "ok": False,
        "pdf_url": '',
        "pdf_pending": True,
        "docx_url": docx_url,
    }

    return JsonResponse({
        "status": "success",
        "message": "Petição de habilitação gerada - Salva em Arquivos.",
        "habilitacao": habilitacao_info,
    })


@login_required
@require_POST
def preview_peticao_custas(request, processo_id=None):
    processo_id = processo_id or request.POST.get('processo_id')
    try:
        processo_id_int = int(processo_id)
    except (TypeError, ValueError):
        return JsonResponse({'custas_preview': ''}, status=400)

    kind = str(request.POST.get('kind') or '').strip().lower()
    if kind not in ('monitoria', 'cobranca'):
        return JsonResponse({'custas_preview': ''}, status=400)

    try:
        processo = get_object_or_404(ProcessoJudicial, pk=processo_id_int)
    except ProcessoJudicial.DoesNotExist:
        return JsonResponse({'custas_preview': ''}, status=404)

    polo_passivo = processo.partes_processuais.filter(tipo_polo='PASSIVO').first()
    if not polo_passivo:
        return JsonResponse({'custas_preview': ''}, status=404)

    contratos_para_monitoria_ids = []
    try:
        posted_json = request.POST.get('contratos_para_monitoria')
        if posted_json:
            contratos_para_monitoria_ids = json.loads(posted_json)
    except (TypeError, json.JSONDecodeError):
        contratos_para_monitoria_ids = []
    contratos_para_monitoria_ids = _parse_contract_ids(contratos_para_monitoria_ids)
    contratos = list(Contrato.objects.filter(id__in=contratos_para_monitoria_ids))

    cnj_entry = _resolve_peticao_cnj_entry(
        processo,
        entry_id=request.POST.get('peticao_cnj_entry_id'),
    )
    processo_override = {}
    if cnj_entry:
        processo_override = {
            'cnj': cnj_entry.cnj or '',
            'uf': cnj_entry.uf or '',
            'vara': cnj_entry.vara or '',
            'tribunal': cnj_entry.tribunal or '',
            'valor_causa': cnj_entry.valor_causa,
        }
    valor_causa_override = _parse_decimal_input(request.POST.get('valor_causa'))
    if valor_causa_override is not None:
        processo_override['valor_causa'] = valor_causa_override

    custas_override = _parse_decimal_input(request.POST.get('custas_total'))
    custas_parcelas = _parse_int_input(request.POST.get('custas_parcelas'))
    custas_valor_parcela = _parse_decimal_input(request.POST.get('custas_valor_parcela'))

    try:
        if kind == 'monitoria':
            docx_bytes = _build_docx_bytes_common(
                processo,
                polo_passivo,
                contratos,
                processo_override=processo_override,
                custas_override=custas_override,
                custas_parcelas=custas_parcelas,
                custas_valor_parcela=custas_valor_parcela,
            )
            start_markers = ["Seja deferido o parcelamento das custas iniciais"]
            end_markers = [
                "Por fim, requer-se",
                "Dá-se à causa",
                "Nestes termos",
            ]
        else:
            total_valor_causa = _get_total_valor_causa(
                contratos,
                processo,
                processo_override.get('valor_causa')
            )
            if custas_override is None:
                custas_override = (total_valor_causa * Decimal('0.025')).quantize(
                    Decimal('0.01'),
                    rounding=ROUND_HALF_UP
                )
            docx_bytes = _build_cobranca_docx_bytes(
                processo,
                polo_passivo,
                contratos,
                processo_override=processo_override,
                custas_override=custas_override,
                custas_parcelas=custas_parcelas,
                custas_valor_parcela=custas_valor_parcela,
            )
            start_markers = ["DAS CUSTAS", "II. DAS CUSTAS"]
            end_markers = ["DOS FATOS"]
        preview_text = _extract_custas_block_from_docx_bytes(
            docx_bytes,
            start_markers,
            end_markers=end_markers
        )
        if kind == 'cobranca' and preview_text:
            preview_text = _complete_cobranca_custas_preview(
                preview_text,
                custas_override,
                parcelas_override=custas_parcelas,
                valor_parcela_override=custas_valor_parcela
            )
    except Exception:
        preview_text = ''

    return JsonResponse({'custas_preview': preview_text})


@login_required
@require_POST
def generate_monitoria_docx_download(request, processo_id=None):
    processo_id = processo_id or request.POST.get('processo_id')
    try:
        processo_id_int = int(processo_id)
    except (TypeError, ValueError):
        return HttpResponse("ID do processo inválido.", status=400)

    try:
        processo = get_object_or_404(ProcessoJudicial, pk=processo_id_int)
        analise = processo.analise_processo
    except ProcessoJudicial.DoesNotExist:
        return HttpResponse("Processo Judicial não encontrado.", status=404)
    except Exception as e:
        logger.error(f"Erro ao buscar análise do processo {processo_id_int}: {e}", exc_info=True)
        return HttpResponse(f"Erro ao buscar dados do processo: {e}", status=500)

    polo_passivo = processo.partes_processuais.filter(tipo_polo='PASSIVO').first()
    if not polo_passivo:
        return HttpResponse("Polo passivo não encontrado para este processo.", status=404)

    contratos_para_monitoria_ids = []
    try:
        posted_json = request.POST.get('contratos_para_monitoria')
        if posted_json:
            contratos_para_monitoria_ids = json.loads(posted_json)
    except (TypeError, json.JSONDecodeError):
        contratos_para_monitoria_ids = []

    if not contratos_para_monitoria_ids:
        contratos_para_monitoria_ids = analise.respostas.get('contratos_para_monitoria', [])

    contratos_monitoria_qs = Contrato.objects.filter(id__in=contratos_para_monitoria_ids).select_related('processo')
    contratos_monitoria = list(contratos_monitoria_qs)
    if not contratos_monitoria:
        return HttpResponse("Nenhum contrato selecionado para monitória na análise deste processo.", status=404)

    try:
        docx_bytes = _build_docx_bytes_common(processo, polo_passivo, contratos_monitoria)
        base_filename = _build_monitoria_base_filename(polo_passivo, contratos_monitoria)
        filename = f"{base_filename}.docx"

        return FileResponse(
            BytesIO(docx_bytes),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            filename=filename,
        )
    except FileNotFoundError as fe:
        return HttpResponse(str(fe), status=500)
    except Exception as e:
        logger.error(f"Erro ao gerar DOCX editável para processo {processo_id}: {e}", exc_info=True)
        return HttpResponse(f"Erro ao gerar DOCX editável: {e}", status=500)


@login_required
@require_GET
def download_monitoria_pdf(request, processo_id=None):
    """
    Download do PDF da monitória com Content-Disposition amigável,
    usando o nome padronizado (substitui underscores por espaços).
    """
    try:
        processo_id_int = int(processo_id or 0)
    except (TypeError, ValueError):
        return HttpResponse("ID do processo inválido.", status=400)

    arquivo_pdf = (
        ProcessoArquivo.objects
        .filter(processo_id=processo_id_int, arquivo__iendswith='.pdf')
        .order_by('-criado_em')
        .first()
    )
    if not arquivo_pdf or not arquivo_pdf.arquivo:
        return HttpResponse("PDF da monitória não encontrado para este processo.", status=404)

    try:
        arquivo_pdf.arquivo.open('rb')
        filename = arquivo_pdf.nome or os.path.basename(arquivo_pdf.arquivo.name)
        # reforça nome legível removendo underscores (storage padrão os adiciona)
        filename = filename.replace('_', ' ')
        return FileResponse(
            arquivo_pdf.arquivo,
            as_attachment=True,
            filename=filename,
            content_type='application/pdf'
        )
    except Exception as exc:
        logger.error("Erro ao preparar download do PDF da monitória: %s", exc, exc_info=True)
        return HttpResponse("Erro ao preparar download do PDF.", status=500)


@login_required
@require_GET
def proxy_arquivo_view(request, arquivo_id):
    """
    Proxy para servir arquivos do S3 com Content-Disposition: inline
    permitindo visualização no iframe em vez de forçar download.
    """
    try:
        arquivo = get_object_or_404(ProcessoArquivo, pk=arquivo_id)
    except ProcessoArquivo.DoesNotExist:
        return HttpResponse("Arquivo não encontrado.", status=404)

    if not arquivo.arquivo:
        return HttpResponse("Arquivo sem conteúdo.", status=404)

    try:
        # Lê o arquivo do S3
        arquivo.arquivo.open('rb')
        file_content = arquivo.arquivo.read()
        arquivo.arquivo.close()

        # Determina o content-type baseado na extensão
        arquivo_name = arquivo.arquivo.name.lower()
        if arquivo_name.endswith('.pdf'):
            content_type = 'application/pdf'
        elif arquivo_name.endswith('.docx'):
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif arquivo_name.endswith('.doc'):
            content_type = 'application/msword'
        elif arquivo_name.endswith('.png'):
            content_type = 'image/png'
        elif arquivo_name.endswith('.jpg') or arquivo_name.endswith('.jpeg'):
            content_type = 'image/jpeg'
        else:
            content_type = 'application/octet-stream'

        # Retorna com headers corretos para visualização inline
        response = HttpResponse(file_content, content_type=content_type)
        response['Content-Disposition'] = 'inline'
        response['Content-Length'] = len(file_content)
        response['X-Content-Type-Options'] = 'nosniff'
        response['Cache-Control'] = 'private, max-age=3600'
        # Headers CORS para permitir PDF.js carregar o PDF
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'GET, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Range'
        return response

    except Exception as exc:
        logger.error("Erro ao servir arquivo via proxy: %s", exc, exc_info=True)
        return HttpResponse("Erro ao carregar arquivo.", status=500)


@login_required
@require_GET
def arquivo_zip_contents_view(request, arquivo_id):
    """Retorna a listagem ordenada dos arquivos internos de um ZIP salvo."""
    try:
        arquivo = get_object_or_404(ProcessoArquivo, pk=arquivo_id)
    except ProcessoArquivo.DoesNotExist:
        return JsonResponse({"status": "error", "error": "Arquivo não encontrado."}, status=404)

    if not arquivo.arquivo:
        return JsonResponse({"status": "error", "error": "Arquivo sem conteúdo."}, status=404)

    try:
        arquivo.arquivo.open('rb')
        zip_bytes = arquivo.arquivo.read()
        arquivo.arquivo.close()
        with zipfile.ZipFile(BytesIO(zip_bytes)) as zip_file:
            entries = [
                info.filename
                for info in zip_file.infolist()
                if info.filename and not info.filename.endswith('/')
            ]
        return JsonResponse({
            "status": "success",
            "file_name": arquivo.nome or os.path.basename(arquivo.arquivo.name),
            "entries": entries,
        })
    except zipfile.BadZipFile:
        return JsonResponse({"status": "error", "error": "O arquivo selecionado não é um ZIP válido."}, status=400)
    except Exception as exc:
        logger.error("Erro ao listar conteúdo do ZIP %s: %s", arquivo_id, exc, exc_info=True)
        return JsonResponse({"status": "error", "error": "Não foi possível listar os arquivos do ZIP."}, status=500)


@login_required
@require_GET
def convert_docx_to_pdf_download(request, arquivo_id):
    """
    Converte um arquivo DOCX existente para PDF on-demand e retorna o download.
    Se o PDF correspondente já existe, retorna ele diretamente.
    """
    logger.info("convert_docx_to_pdf_download chamado: arquivo_id=%s", arquivo_id)
    try:
        arquivo = get_object_or_404(ProcessoArquivo, pk=arquivo_id)
    except ProcessoArquivo.DoesNotExist:
        logger.error("Arquivo %s não encontrado", arquivo_id)
        return HttpResponse("Arquivo não encontrado.", status=404)

    arquivo_path = arquivo.arquivo.name if arquivo.arquivo else ''
    if not arquivo_path:
        logger.error("Arquivo %s sem conteúdo", arquivo_id)
        return HttpResponse("Arquivo sem conteúdo.", status=404)

    logger.info("Arquivo path: %s", arquivo_path)

    # Verifica se é um DOCX
    is_docx = arquivo_path.lower().endswith('.docx')
    is_pdf = arquivo_path.lower().endswith('.pdf')

    # Por padrão visualiza inline, só baixa se passar ?download=1
    download = request.GET.get('download') == '1'
    if is_pdf:
        # Já é PDF, retorna direto
        try:
            arquivo.arquivo.open('rb')
            pdf_content = arquivo.arquivo.read()
            arquivo.arquivo.close()

            filename = (arquivo.nome or os.path.basename(arquivo_path)).replace('_', ' ')

            if download:
                # Força download
                return FileResponse(
                    BytesIO(pdf_content),
                    as_attachment=True,
                    filename=filename,
                    content_type='application/pdf'
                )
            else:
                # Visualiza inline (sem filename para evitar download)
                response = HttpResponse(pdf_content, content_type='application/pdf')
                response['Content-Disposition'] = 'inline'
                response['Content-Length'] = len(pdf_content)
                response['X-Content-Type-Options'] = 'nosniff'
                response['Cache-Control'] = 'private, max-age=3600'
                # Headers CORS para permitir PDF.js carregar o PDF
                response['Access-Control-Allow-Origin'] = '*'
                response['Access-Control-Allow-Methods'] = 'GET, OPTIONS'
                response['Access-Control-Allow-Headers'] = 'Range'
                return response
        except Exception as exc:
            logger.error("Erro ao acessar PDF: %s", exc, exc_info=True)
            return HttpResponse("Erro ao acessar PDF.", status=500)

    if not is_docx:
        return HttpResponse("Arquivo não é DOCX nem PDF.", status=400)

    # Verifica se já existe um PDF correspondente
    base_name = arquivo_path.rsplit('.', 1)[0]
    pdf_path = base_name + '.pdf'

    # Tenta encontrar o PDF existente no mesmo processo
    existing_pdf = ProcessoArquivo.objects.filter(
        processo=arquivo.processo,
        arquivo__iendswith='.pdf',
        nome__icontains=arquivo.nome.rsplit('.', 1)[0] if arquivo.nome else ''
    ).first()

    if existing_pdf and existing_pdf.arquivo:
        try:
            existing_pdf.arquivo.open('rb')
            pdf_content = existing_pdf.arquivo.read()
            existing_pdf.arquivo.close()

            filename = (existing_pdf.nome or os.path.basename(existing_pdf.arquivo.name)).replace('_', ' ')

            if download:
                return FileResponse(
                    BytesIO(pdf_content),
                    as_attachment=True,
                    filename=filename,
                    content_type='application/pdf'
                )
            else:
                # Visualiza inline
                response = HttpResponse(pdf_content, content_type='application/pdf')
                response['Content-Disposition'] = 'inline'
                response['Content-Length'] = len(pdf_content)
                response['X-Content-Type-Options'] = 'nosniff'
                response['Cache-Control'] = 'private, max-age=3600'
                # Headers CORS para permitir PDF.js carregar o PDF
                response['Access-Control-Allow-Origin'] = '*'
                response['Access-Control-Allow-Methods'] = 'GET, OPTIONS'
                response['Access-Control-Allow-Headers'] = 'Range'
                return response
        except Exception:
            pass  # Se falhar, tenta converter

    # Converte DOCX para PDF
    try:
        arquivo.arquivo.open('rb')
        docx_bytes = arquivo.arquivo.read()
        arquivo.arquivo.close()
    except Exception as exc:
        logger.error("Erro ao ler DOCX para conversão: %s", exc, exc_info=True)
        return HttpResponse("Erro ao ler arquivo DOCX.", status=500)

    logger.info("Iniciando conversão DOCX para PDF (tamanho: %d bytes)", len(docx_bytes))
    pdf_bytes = _convert_docx_to_pdf_bytes(docx_bytes)
    if not pdf_bytes:
        logger.error("Conversão falhou: pdf_bytes é None ou vazio")
        return HttpResponse(
            "Não foi possível converter o DOCX para PDF. "
            "O conversor não está disponível no servidor.",
            status=500
        )

    logger.info("Conversão bem-sucedida: PDF com %d bytes", len(pdf_bytes))

    # Retorna o PDF para visualização ou download
    pdf_filename = (arquivo.nome or os.path.basename(arquivo_path)).rsplit('.', 1)[0] + '.pdf'

    if download:
        return FileResponse(
            BytesIO(pdf_bytes),
            as_attachment=True,
            filename=pdf_filename.replace('_', ' '),
            content_type='application/pdf'
        )
    else:
        # Visualiza inline
        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = 'inline'
        response['Content-Length'] = len(pdf_bytes)
        response['X-Content-Type-Options'] = 'nosniff'
        response['Cache-Control'] = 'private, max-age=3600'
        # Headers CORS para permitir PDF.js carregar o PDF
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'GET, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Range'
        return response


def _convert_pdf_to_docx_bytes(pdf_bytes):
    """
    Converte PDF para DOCX usando pdf2docx (100% Python).
    """
    try:
        from pdf2docx import Converter

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            pdf_path = tmpdir_path / "input.pdf"
            docx_path = tmpdir_path / "output.docx"

            pdf_path.write_bytes(pdf_bytes)

            # Converte PDF para DOCX
            cv = Converter(str(pdf_path))
            cv.convert(str(docx_path))
            cv.close()

            if docx_path.exists():
                return docx_path.read_bytes()

    except ImportError as e:
        logger.error("Biblioteca pdf2docx não instalada: %s", e)
    except Exception as exc:
        logger.error("Erro ao converter PDF para DOCX: %s", exc, exc_info=True)

    return None


@login_required
@require_GET
def convert_pdf_to_docx_download(request, arquivo_id):
    """
    Converte um arquivo PDF existente para DOCX on-demand e retorna o download.
    """
    try:
        arquivo = get_object_or_404(ProcessoArquivo, pk=arquivo_id)
    except ProcessoArquivo.DoesNotExist:
        return HttpResponse("Arquivo não encontrado.", status=404)

    arquivo_path = arquivo.arquivo.name if arquivo.arquivo else ''
    if not arquivo_path:
        return HttpResponse("Arquivo sem conteúdo.", status=404)

    is_pdf = arquivo_path.lower().endswith('.pdf')
    is_docx = arquivo_path.lower().endswith('.docx')

    if is_docx:
        # Já é DOCX, retorna direto
        try:
            arquivo.arquivo.open('rb')
            filename = (arquivo.nome or os.path.basename(arquivo_path)).replace('_', ' ')
            return FileResponse(
                arquivo.arquivo,
                as_attachment=True,
                filename=filename,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except Exception as exc:
            logger.error("Erro ao baixar DOCX: %s", exc, exc_info=True)
            return HttpResponse("Erro ao baixar DOCX.", status=500)

    if not is_pdf:
        return HttpResponse("Arquivo não é PDF nem DOCX.", status=400)

    # Tenta encontrar o DOCX existente no mesmo processo
    existing_docx = ProcessoArquivo.objects.filter(
        processo=arquivo.processo,
        arquivo__iendswith='.docx',
        nome__icontains=arquivo.nome.rsplit('.', 1)[0] if arquivo.nome else ''
    ).first()

    if existing_docx and existing_docx.arquivo:
        try:
            existing_docx.arquivo.open('rb')
            filename = (existing_docx.nome or os.path.basename(existing_docx.arquivo.name)).replace('_', ' ')
            return FileResponse(
                existing_docx.arquivo,
                as_attachment=True,
                filename=filename,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except Exception:
            pass  # Se falhar, tenta converter

    # Converte PDF para DOCX
    try:
        arquivo.arquivo.open('rb')
        pdf_bytes = arquivo.arquivo.read()
        arquivo.arquivo.close()
    except Exception as exc:
        logger.error("Erro ao ler PDF para conversão: %s", exc, exc_info=True)
        return HttpResponse("Erro ao ler arquivo PDF.", status=500)

    docx_bytes = _convert_pdf_to_docx_bytes(pdf_bytes)
    if not docx_bytes:
        return HttpResponse(
            "Não foi possível converter o PDF para DOCX.",
            status=500
        )

    # Retorna o DOCX para download (sem salvar em Arquivos)
    docx_filename = (arquivo.nome or os.path.basename(arquivo_path)).rsplit('.', 1)[0] + '.docx'
    return FileResponse(
        BytesIO(docx_bytes),
        as_attachment=True,
        filename=docx_filename.replace('_', ' '),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
